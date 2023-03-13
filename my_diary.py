from customtkinter import * 
from pickle import * 
from tkinter import filedialog
from docx.shared import RGBColor
import os 
import sys 
import docx 

class Programm(CTk):
	def __init__(self):
		super().__init__()

		set_appearance_mode("Dark")
		set_default_color_theme("green")

		self.main_screen = CTk()
		self.main_screen.geometry("1920x1080")
		self.main_screen.title("my_diary")
		self.main_screen.iconbitmap("my_diary_icon.ico")

		self.main_screen.bind('<F1>', self.__programm_version__) 
		self.main_screen.bind('<Control_L>' + '<t>', self.__theme__) 
		self.main_screen.bind('<Control_L>' + '<f>', self.__fullscreen__) 
		self.main_screen.bind('<Escape>' + '<e>', self.__exit__)

		self.main_screen_theme_numbers = 1
		self.main_screen_fullscreen_numbers = 1

		with open("my_diary_settings.pickle", "rb") as self.data:
			self.new_data = load(self.data)

		with open('my_diary_saved_text.pickle', "rb") as self.text_data:
			self.new_text_data = load(self.text_data)

		with open("my_diary_autosave_settings.pickle", "rb") as self.autosave_data:
			self.new_autosaved_data = load(self.autosave_data)

		self.main_screen_save_button = CTkButton(master=self.main_screen, text="сачувај текст", height=20, width=25, corner_radius=0, command=self.__save__)
		self.main_screen_save_button.grid(column=0, row=0)

		self.main_screen_save_as_txt_button = CTkButton(master=self.main_screen, text="сачувај као txt", height=20, width=25, corner_radius=0, command=self.__save_text_as_txt__)

		self.main_screen_save_as_docx_button = CTkButton(master=self.main_screen, text="сачувај као docx", height=20, width=25, corner_radius=0, command=self.__save_as_docx__)

		self.main_screen_save_as_py_button = CTkButton(master=self.main_screen, text="сачувај као python", height=20, width=25, corner_radius=0, command=self.__save_text_as_py__)

		self.main_screen_save_as_java_button = CTkButton(master=self.main_screen, text="сачувај као java", height=20, width=25, corner_radius=0, command=self.__save_text_as_java__)

		self.main_screen_save_as_kotlin_button = CTkButton(master=self.main_screen, text="сачувај као kotlin", height=20, width=25, corner_radius=0, command=self.__save_text_as_kotlin__)

		self.main_screen_save_button_clicks = 1

		self.main_screen_clear_button = CTkButton(master=self.main_screen, text="обриши текст", height=20, width=25, corner_radius=0, command=self.__clear_text__)
		self.main_screen_clear_button.grid(column=1, row=0)

		self.main_screen_edit_text_button = CTkButton(master=self.main_screen, text="уреди текст", height=20, width=25, corner_radius=0, command=self.__edit_text_)
		self.main_screen_edit_text_button.grid(column=2, row=0)

		self.main_screen_edit_font_button = CTkComboBox(master=self.main_screen, height=20, width=75, corner_radius=0, values=["System", "Terminal", "Ubuntu", "Script", "Roman", "Modern", "MS Serif"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", dropdown_hover_color="white", text_color="white", command=self.__edit_text_font_and_size__)

		self.main_screen_edit_size_button = CTkComboBox(master=self.main_screen, height=20, width=75, corner_radius=0, values=["1", "2", "4", "5", "6", "8", "11", "12", "13", "14", "16", "18", "20", "22", "24", "26", "28", "30", "32", "34", "36", "38", "40", "42", "44", "46", "48", "50", "60", "70", "80", "90", "100"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", dropdown_hover_color="white", text_color="white", command=self.__edit_text_font_and_size__)

		self.main_screen_edit_color_button = CTkComboBox(master=self.main_screen, height=20, width=75, corner_radius=0, values=["black", "white", "red", "green", "blue"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", dropdown_hover_color="white", text_color="white", command=self.__edit_text_color__)

		self.main_screen_edit_clicks = 1

		self.main_screen_open_file_button = CTkButton(master=self.main_screen, text="отвори фајл", height=20, width=25, corner_radius=0, command=self.__open_file__)
		self.main_screen_open_file_button.grid(column=3, row=0)
	
		self.main_screen_open_file_txt_button = CTkButton(master=self.main_screen, text="отвори txt фајл", height=20, width=25, corner_radius=0, command=self.__open_file_txt__)
	
		self.main_screen_open_file_docx_button = CTkButton(master=self.main_screen, text="отвори docx фајл", height=20, width=25, corner_radius=0, command=self.__open_file_docx__)
	
		self.main_screen_open_file_py_button = CTkButton(master=self.main_screen, text="отвори python фајл", height=20, width=25, corner_radius=0, command=self.__open_file_py__)

		self.main_screen_open_file_clicks = 1

		self.main_screen_coding_button = CTkButton(master=self.main_screen, text="кодирање", height=20, width=25, corner_radius=0, command=self.__coding__)
		self.main_screen_coding_button.grid(column=4, row=0)

		self.main_screen_coding_binar_values = {"a": "0", "b": "1", "c": "10", "d": "11", "e": "100", "f": "101", "g": "110", "h": "111", "i": "1000", "j": "1001", "k": "1010", "l": "1011", "m": "1100", "n": "1101", "o": "1110", "p": "1111", "q": "10000", "r": "10001", "s": "10010", "t": "10011", "u": "10100", "v": "10101", "w": "10111", "x": "11000", "y": "11001", "z": "11011", " ": " ", ".": "."}

		self.main_screen_coding_ceaser_values = {"a": "d", "b": "e", "c": "f", "d": "g", "e": "h", "f": "i", "g": "j", "h": "k", "i": "l", "j": "m", "k": "n", "l": "o", "m": "p", "n": "q", "o": "r", "p": "s", "q": "t", "r": "u", "s": "v", "t": "w", "u": "x", "v": "y", "w": "z", "x": "a", "y": "b", "z": "c", " ": " ", ".": "."}
		
		self.main_screen_coding_fibonachi_values = {"a": "0", "b": "1", "c": "2", "d": "3", "e": "5", "f": "8", "g": "13", "h": "21", "i": "34", "j": "55", "k": "89", "l": "144", "m": "233", "n": "377", "o": "610", "p": "987", "q": "1597", "r": "2584", "s": "4181", "t": "6765", "u": "10946", "v": "17711", "w": "28657", "x": "46368", "y": "75025", "z": "121393", " ": " ", ".": "."}

		self.main_screen_coding_numerical_values = {"a": "1", "b": "2", "c" : "3", "d": "4", "e": "5", "f": "6", "g": "7", "h": "8", "i": "9", "j": "10", "k": "11", "l": "12", "m": "13", "q": "14", "r": "15", "s": "16", "t": "17", "u": "18", "v": "19", "w": "20", "x": "21", "y": "22", "z": "23", " ": " ", ".": "."}

		self.main_screen_coding_mode_button = CTkComboBox(master=self.main_screen, height=20, width=75, corner_radius=0, values=["binar", "ceaser", "fibonachi", "numerical"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_hover_color="white")

		self.main_screen_coding_function_button = CTkButton(master=self.main_screen, text="кодирај", height=20, width=25, corner_radius=0, command=self.__message_encoding__)

		self.main_screen_coding_clicks = 1

		self.main_screen_settings_button = CTkButton(master=self.main_screen, text="подешавања", height=20, width=25, corner_radius=0, command=self.__settings__)
		self.main_screen_settings_button.grid(column=5, row=0)

		self.main_screen_settings_text = CTkLabel(master=self.main_screen, text="подешавања", font=("Roboto Bold", 72))

		self.main_screen_settings_language_text = CTkLabel(master=self.main_screen, text="језици", font=("Roboto Bold", 36))

		self.main_screen_settings_language_option = CTkSegmentedButton(master=self.main_screen, values=["Србски", "English", "Русский", "Deutch", "Español"], command=self.__language_settings__)

		self.language_data = str(self.new_data)

		self.main_screen_settings_language_option.set(self.language_data)
		
		if self.language_data == "Србски":
			self.main_screen_save_button.configure(text="сачувај текст")
			self.main_screen_save_as_txt_button.configure(text="сачувај као txt")
			self.main_screen_save_as_docx_button.configure(text="сачувај као docx")
			self.main_screen_save_as_py_button.configure(text="сачувај као python")
			self.main_screen_save_as_java_button.configure(text="сачувај као java")
			self.main_screen_save_as_kotlin_button.configure(text="сачувај као kotlin")

			self.main_screen_clear_button.configure(text="обриши текст")
			self.main_screen_edit_text_button.configure(text="уреди текст")
			
			self.main_screen_open_file_button.configure(text="отвори фајл")
			self.main_screen_open_file_txt_button.configure(text="отвори txt фајл")
			self.main_screen_open_file_docx_button.configure(text="отвори docx фајл")
			self.main_screen_open_file_py_button.configure(text="отвори python фајл")

			self.main_screen_coding_button.configure(text="кодирање")
			self.main_screen_coding_function_button.configure(text="кодирај")

			self.main_screen_settings_button.configure(text="подешавања")
			self.main_screen_settings_text.configure(text="подешавања")
			self.main_screen_settings_language_text.configure(text="језици")

		elif self.language_data == "English":
			self.main_screen_save_button.configure(text="save text")
			self.main_screen_save_as_txt_button.configure(text="save as txt")
			self.main_screen_save_as_docx_button.configure(text="save as docx")
			self.main_screen_save_as_py_button.configure(text="save as python")
			self.main_screen_save_as_java_button.configure(text="save as java")
			self.main_screen_save_as_kotlin_button.configure(text="save as kotlin")

			self.main_screen_clear_button.configure(text="clear text")
			self.main_screen_edit_text_button.configure(text="edit text")
			
			self.main_screen_open_file_button.configure(text="open file")
			self.main_screen_open_file_txt_button.configure(text="open txt file")
			self.main_screen_open_file_docx_button.configure(text="open docx file")
			self.main_screen_open_file_py_button.configure(text="open python file")

			self.main_screen_coding_button.configure(text="coding")
			self.main_screen_coding_function_button.configure(text="code")

			self.main_screen_settings_button.configure(text="settings")
			self.main_screen_settings_text.configure(text="settings")
			self.main_screen_settings_language_text.configure(text="languages")

		elif self.language_data == "Русский":
			self.main_screen_save_button.configure(text="сохранить текст")
			self.main_screen_save_as_txt_button.configure(text="сохранить как текст")
			self.main_screen_save_as_docx_button.configure(text="сохранить как ворд")
			self.main_screen_save_as_py_button.configure(text="сохранить как python")
			self.main_screen_save_as_java_button.configure(text="сохранить как java")
			self.main_screen_save_as_kotlin_button.configure(text="сохранить как kotlin")

			self.main_screen_clear_button.configure(text="удалить текст")
			self.main_screen_edit_text_button.configure(text="редактировать текст")
			
			self.main_screen_open_file_button.configure(text="открыть файл")
			self.main_screen_open_file_txt_button.configure(text="открыть текстовый файл")
			self.main_screen_open_file_docx_button.configure(text="открыть ворд файл")
			self.main_screen_open_file_py_button.configure(text="открыть python файл")
			
			self.main_screen_coding_button.configure(text="кодировка")
			self.main_screen_coding_function_button.configure(text="кодировать")

			self.main_screen_settings_button.configure(text="настройки")
			self.main_screen_settings_text.configure(text="настройки")
			self.main_screen_settings_language_text.configure(text="языки")

		elif self.language_data == "Deutch":
			self.main_screen_save_button.configure(text="Texte speichern")
			self.main_screen_save_as_txt_button.configure(text="als txt speichern")
			self.main_screen_save_as_docx_button.configure(text="als doсx speichern")
			self.main_screen_save_as_py_button.configure(text="als Python speichern")
			self.main_screen_save_as_java_button.configure(text="als Java speichern")
			self.main_screen_save_as_kotlin_button.configure(text="als Kotlin speichern")

			self.main_screen_clear_button.configure(text="Texte löschen")
			self.main_screen_edit_text_button.configure(text="Text bearbeiten")
			
			self.main_screen_open_file_button.configure(text="Datei öffnen")
			self.main_screen_open_file_txt_button.configure(text="Textdatei öffnen")
			self.main_screen_open_file_docx_button.configure(text="Word-Datei öffnen")
			self.main_screen_open_file_py_button.configure(text="Python-Datei öffnen")
			
			self.main_screen_coding_button.configure(text="Codierung")
			self.main_screen_coding_function_button.configure(text="kodieren")

			self.main_screen_settings_button.configure(text="Einstellungen")
			self.main_screen_settings_text.configure(text="Einstellungen")
			self.main_screen_settings_language_text.configure(text="Sprachen")

		elif self.language_data == "Español":
			self.main_screen_save_button.configure(text="guardar texto")
			self.main_screen_save_as_txt_button.configure(text="guardar como txt")
			self.main_screen_save_as_docx_button.configure(text="guardar como docx")
			self.main_screen_save_as_py_button.configure(text="guardar como python")
			self.main_screen_save_as_java_button.configure(text="guardar como java")
			self.main_screen_save_as_kotlin_button.configure(text="guardar como kotlin")

			self.main_screen_clear_button.configure(text="borrar texto")
			self.main_screen_edit_text_button.configure(text="editar texto")
			
			self.main_screen_open_file_button.configure(text="abrir documento")
			self.main_screen_open_file_txt_button.configure(text="abrir archivo de texto")
			self.main_screen_open_file_docx_button.configure(text="abrir archivo de palabra")
			self.main_screen_open_file_py_button.configure(text="abrir archivo python")
			
			self.main_screen_coding_button.configure(text="codificacióng")
			self.main_screen_coding_function_button.configure(text="codificar")

			self.main_screen_settings_button.configure(text="ajustes")
			self.main_screen_settings_text.configure(text="ajustes")
			self.main_screen_settings_language_text.configure(text="lenguas")

		self.main_screen_settings_autosave_text = CTkLabel(master=self.main_screen, text="autosave", font=("Roboto Bold", 36))

		self.main_screen_settings_autosave_value = StringVar(value=self.new_autosaved_data)

		self.main_screen_settings_autosave_switch_varuable = str(self.new_autosaved_data)

		self.main_screen_settings_autosave_switch = CTkSwitch(master=self.main_screen, text="autosave", variable=self.main_screen_settings_autosave_value,onvalue="on", offvalue="off", command=self.__autosave_settings__)
			
		self.main_screen_settings_button_clicks = 1

		self.main_screen_frame = CTkFrame(master=self.main_screen, height=776, width=1535, border_width=2, border_color=("black", "white"), corner_radius=0)
		self.main_screen_frame.place(x=0, y=21)

		self.main_screen_frame_textbox = CTkTextbox(master=self.main_screen_frame, height=770, width=1529, corner_radius=0)
		self.main_screen_frame_textbox.place(x=3, y=3)

		self.autosaved_text = str(self.new_text_data)

		self.main_screen_settings_autosave_switch_value = self.main_screen_settings_autosave_value.get()

		if self.main_screen_settings_autosave_switch_varuable == "on":
			self.main_screen_frame_textbox.insert("1.0", self.autosaved_text)

			self.main_screen.bind('<Return>', self.__text_autosave__) 
			self.main_screen.bind('<Control_L>' + '<s>', self.__text_autosave__)

		elif self.main_screen_settings_autosave_switch_varuable == "off":
			self.main_screen.unbind('<Return>')
			self.main_screen.unbind('<Control_L>' + '<s>')
				  
		self.main_screen_coding_frame = CTkFrame(master=self.main_screen, height=776, width=717.5, border_width=2, border_color=("black", "white"), corner_radius=0)

		self.main_screen_coding_frame_textbox = CTkTextbox(master=self.main_screen_coding_frame, height=770, width=709.5, corner_radius=0)

	def __run__(self):
		self.main_screen.mainloop()

	def __clear_text__(self):
		self.main_screen_frame_textbox.delete("1.0", END)

	def __save__(self):
		self.main_screen_save_button_clicks += 1
		if self.main_screen_save_button_clicks % 2 == 0:
			self.main_screen_save_as_txt_button.grid(column=1, row=0)
			self.main_screen_save_as_docx_button.grid(column=2, row=0)
			self.main_screen_save_as_py_button.grid(column=3, row=0)
			self.main_screen_save_as_java_button.grid(column=4, row=0)
			self.main_screen_save_as_kotlin_button.grid(column=5, row=0)

			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_coding_button.grid_forget()     
			self.main_screen_settings_button.grid_forget()       

		else:
			self.main_screen_clear_button.grid(column=1, row=0)
			self.main_screen_edit_text_button.grid(column=2, row=0)
			self.main_screen_open_file_button.grid(column=3, row=0)
			self.main_screen_coding_button.grid(column=4, row=0)
			self.main_screen_settings_button.grid(column=5, row=0)

			self.main_screen_save_as_txt_button.grid_forget()
			self.main_screen_save_as_docx_button.grid_forget()
			self.main_screen_save_as_py_button.grid_forget()
			self.main_screen_save_as_java_button.grid_forget()
			self.main_screen_save_as_kotlin_button.grid_forget()
			
	def __save_text_as_txt__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_as_docx__(self):
		self.file = docx.Document()
		self.file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file_run = self.file.add_paragraph().add_run(self.file_data)
		self.font = self.file_run.font
		if self.main_screen_edit_color_button_data == "black":
			self.font.color.rgb = RGBColor(0, 0, 0)
		
		elif self.main_screen_edit_color_button_data == "white":
			self.font.color.rgb = RGBColor(255, 255, 255)

		elif self.main_screen_edit_color_button_data == "red":
			self.font.color.rgb = RGBColor(250, 0, 0)

		elif self.main_screen_edit_color_button_data == "green":
			self.font.color.rgb = RGBColor(0, 255, 0)

		elif self.main_screen_edit_color_button_data == "blue":
			self.font.color.rgb = RGBColor(0, 0, 255)
			
		self.file.save(filedialog.asksaveasfilename(filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")]))

	def __save_text_as_py__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Python file (*.py)", "*.py")], defaultextension=[("Python file (*.py)", "*.py")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_java__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Java file (*.java)", "*.java")], defaultextension=[("Java file (*.java)", "*.java")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_kotlin__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Kotlin file (*.kt)", "*.kt")], defaultextension=[("Kotlin file (*.kt)", "*.kt")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __edit_text_(self):
		self.main_screen_edit_clicks += 1
		if self.main_screen_edit_clicks % 2 == 0:
			self.main_screen_edit_text_button.grid(column=0, row=0)
			self.main_screen_edit_font_button.grid(column=1, row=0)
			self.main_screen_edit_size_button.grid(column=2, row=0)
			self.main_screen_edit_color_button.grid(column=3, row=0)

			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_coding_button.grid_forget()
			self.main_screen_settings_button.grid_forget()

		else:
			self.main_screen_save_button.grid(column=0, row=0)
			self.main_screen_clear_button.grid(column=1, row=0)
			self.main_screen_edit_text_button.grid(column=2, row=0)
			self.main_screen_open_file_button.grid(column=3, row=0)
			self.main_screen_coding_button.grid(column=4, row=0)
			self.main_screen_settings_button.grid(column=5, row=0)

			self.main_screen_edit_font_button.grid_forget()
			self.main_screen_edit_size_button.grid_forget()
			self.main_screen_edit_color_button.grid_forget()

	def __edit_text_font_and_size__(self, configure):
		self.main_screen_edit_font_button_data = self.main_screen_edit_font_button.get()
		self.main_screen_edit_size_button_data = self.main_screen_edit_size_button.get()
		self.main_screen_frame_textbox.configure(font=(self.main_screen_edit_font_button_data, float(self.main_screen_edit_size_button_data)))

	def __edit_text_color__(self, configure):
		self.main_screen_edit_color_button_data = self.main_screen_edit_color_button.get()
		self.main_screen_frame_textbox.configure(text_color=self.main_screen_edit_color_button_data)

	def __open_file__(self):
		self.main_screen_open_file_clicks += 1
		if self.main_screen_open_file_clicks % 2 == 0:
			self.main_screen_open_file_button.grid(column=0, row=0)
			self.main_screen_open_file_txt_button.grid(column=1, row=0)
			self.main_screen_open_file_docx_button.grid(column=2, row=0)
			self.main_screen_open_file_py_button.grid(column=3, row=0)

			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_coding_button.grid_forget()
			self.main_screen_settings_button.grid_forget()

		else:
			self.main_screen_save_button.grid(column=0, row=0)
			self.main_screen_clear_button.grid(column=1, row=0)
			self.main_screen_edit_text_button.grid(column=2, row=0)
			self.main_screen_open_file_button.grid(column=3, row=0)
			self.main_screen_coding_button.grid(column=4, row=0)
			self.main_screen_settings_button.grid(column=5, row=0)

			self.main_screen_open_file_txt_button.grid_forget()
			self.main_screen_open_file_docx_button.grid_forget()
			self.main_screen_open_file_py_button.grid_forget()

	def __open_file_txt__(self):
		self.openned_file = open(filedialog.askopenfilename(title="open txt file"), "r+")
		self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

	def __open_file_docx__(self):
		self.openned_file = docx.Document(filedialog.askopenfilename(title="open docx file"))
		self.openned_file_data = []
		for self.paragraphs in self.openned_file.paragraphs:
			self.openned_file_data.append(self.paragraphs.text)

		self.main_screen_frame_textbox.insert("1.0", "\n".join(self.openned_file_data))

	def __open_file_py__(self):
		self.openned_file = open(filedialog.askopenfilename(title="open file"), "r+", encoding="UTF-8")
		self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

	def __coding__(self):
		self.main_screen_coding_clicks += 1
		if self.main_screen_coding_clicks % 2 == 0:
			self.main_screen_frame.configure(width=715.5)
			self.main_screen_frame_textbox.configure(width=709.5)

			self.main_screen_coding_button.grid(column=0, row=0)
			self.main_screen_coding_frame.place(x=729, y=21)
			self.main_screen_coding_frame_textbox.place(x=3, y=3)
			self.main_screen_coding_mode_button.place(x=729, y=0)
			self.main_screen_coding_function_button.place(x=803.5, y=0)

			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_settings_button.grid_forget()

		else:
			self.main_screen_frame.configure(width=1535)
			self.main_screen_frame_textbox.configure(width=1529)

			self.main_screen_coding_frame.place_forget()
			self.main_screen_coding_frame_textbox.place_forget()
			self.main_screen_coding_mode_button.place_forget()
			self.main_screen_coding_function_button.place_forget()

			self.main_screen_save_button.grid(column=0, row=0)
			self.main_screen_clear_button.grid(column=1, row=0)
			self.main_screen_edit_text_button.grid(column=2, row=0)
			self.main_screen_open_file_button.grid(column=3, row=0)
			self.main_screen_coding_button.grid(column=4, row=0)
			self.main_screen_settings_button.grid(column=5, row=0)

			self.main_screen_coding_frame_textbox.delete("1.0", END)

	def __message_encoding__(self):
		self.main_screen_coding_frame_textbox.delete("1.0", END)

		self.main_screen_coding_mode_button_data = self.main_screen_coding_mode_button.get()
		self.main_screen_frame_textbox_data = self.main_screen_frame_textbox.get("1.0", END)
		if self.main_screen_coding_mode_button_data == "binar":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_binar_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_binar_values.get(key), "-1.0")
		
		elif self.main_screen_coding_mode_button_data == "ceaser":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_ceaser_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_ceaser_values.get(key), "-1.0")

		elif self.main_screen_coding_mode_button_data == "fibonachi":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_fibonachi_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_fibonachi_values.get(key), "-1.0")

		elif self.main_screen_coding_mode_button_data == "numerical":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_numerical_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_numerical_values.get(key), "-1.0")

	def __settings__(self):
		self.main_screen_settings_button_clicks += 1
		if self.main_screen_settings_button_clicks % 2 == 0:
			self.main_screen_settings_button.grid(column=0, row=0)
			self.main_screen_settings_text.place(x=3, y=21)
			self.main_screen_settings_language_text.place(x=3, y=102)
			self.main_screen_settings_language_option.place(x=18, y=142)
			self.main_screen_settings_autosave_text.place(x=3, y=182)
			self.main_screen_settings_autosave_switch.place(x=18, y=222)

			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_coding_button.grid_forget()
			self.main_screen_frame.place_forget()
	   
		else:
			self.main_screen_save_button.grid(column=0, row=0)
			self.main_screen_clear_button.grid(column=1, row=0)
			self.main_screen_edit_text_button.grid(column=2, row=0)
			self.main_screen_open_file_button.grid(column=3, row=0)
			self.main_screen_coding_button.grid(column=4, row=0)
			self.main_screen_settings_button.grid(column=5, row=0)
			self.main_screen_frame.place(x=0, y=21)

			self.main_screen_settings_text.place_forget()
			self.main_screen_settings_language_text.place_forget()
			self.main_screen_settings_language_option.place_forget()
			self.main_screen_settings_autosave_text.place_forget()
			self.main_screen_settings_autosave_switch.place_forget()

	def __language_settings__(self, pickle):
		self.main_screen_settings_language_option_varuable = self.main_screen_settings_language_option.get()
		if self.main_screen_settings_language_option_varuable == "Србски":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_varuable, self.data)

			self.main_screen_save_button.configure(text="сачувај текст")
			self.main_screen_save_as_txt_button.configure(text="сачувај као txt")
			self.main_screen_save_as_docx_button.configure(text="сачувај као docx")
			self.main_screen_save_as_py_button.configure(text="сачувај као python")
			self.main_screen_save_as_java_button.configure(text="сачувај као java")
			self.main_screen_save_as_kotlin_button.configure(text="сачувај као kotlin")

			self.main_screen_clear_button.configure(text="обриши текст")
			self.main_screen_edit_text_button.configure(text="уреди текст")
			
			self.main_screen_open_file_button.configure(text="отвори фајл")
			self.main_screen_open_file_txt_button.configure(text="отвори txt фајл")
			self.main_screen_open_file_docx_button.configure(text="отвори docx фајл")
			self.main_screen_open_file_py_button.configure(text="отвори python фајл")

			self.main_screen_coding_button.configure(text="кодирање")
			self.main_screen_coding_function_button.configure(text="кодирај")

			self.main_screen_settings_button.configure(text="подешавања")
			self.main_screen_settings_text.configure(text="подешавања")
			self.main_screen_settings_language_text.configure(text="језици")

		elif self.main_screen_settings_language_option_varuable == "English":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_varuable, self.data)

			self.main_screen_save_button.configure(text="save text")
			self.main_screen_save_as_txt_button.configure(text="save as txt")
			self.main_screen_save_as_docx_button.configure(text="save as docx")
			self.main_screen_save_as_py_button.configure(text="save as python")
			self.main_screen_save_as_java_button.configure(text="save as java")
			self.main_screen_save_as_kotlin_button.configure(text="save as kotlin")

			self.main_screen_clear_button.configure(text="clear text")
			self.main_screen_edit_text_button.configure(text="edit text")
			
			self.main_screen_open_file_button.configure(text="open file")
			self.main_screen_open_file_txt_button.configure(text="open txt file")
			self.main_screen_open_file_docx_button.configure(text="open docx file")
			self.main_screen_open_file_py_button.configure(text="open python file")

			self.main_screen_coding_button.configure(text="coding")
			self.main_screen_coding_function_button.configure(text="code")

			self.main_screen_settings_button.configure(text="settings")
			self.main_screen_settings_text.configure(text="settings")
			self.main_screen_settings_language_text.configure(text="languages")

		elif self.main_screen_settings_language_option_varuable == "Русский":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_varuable, self.data)

			self.main_screen_save_button.configure(text="сохранить текст")
			self.main_screen_save_as_txt_button.configure(text="сохранить как текст")
			self.main_screen_save_as_docx_button.configure(text="сохранить как ворд")
			self.main_screen_save_as_py_button.configure(text="сохранить как python")
			self.main_screen_save_as_java_button.configure(text="сохранить как java")
			self.main_screen_save_as_kotlin_button.configure(text="сохранить как kotlin")

			self.main_screen_clear_button.configure(text="удалить текст")
			self.main_screen_edit_text_button.configure(text="редактировать текст")
			
			self.main_screen_open_file_button.configure(text="открыть файл")
			self.main_screen_open_file_txt_button.configure(text="открыть текстовый файл")
			self.main_screen_open_file_docx_button.configure(text="открыть ворд файл")
			self.main_screen_open_file_py_button.configure(text="открыть python файл")
			
			self.main_screen_coding_button.configure(text="кодировка")
			self.main_screen_coding_function_button.configure(text="кодировать")

			self.main_screen_settings_button.configure(text="настройки")
			self.main_screen_settings_text.configure(text="настройки")
			self.main_screen_settings_language_text.configure(text="языки")

		elif self.main_screen_settings_language_option_varuable == "Deutch":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_varuable, self.data)

			self.main_screen_save_button.configure(text="Texte speichern")
			self.main_screen_save_as_txt_button.configure(text="als txt speichern")
			self.main_screen_save_as_docx_button.configure(text="als doсx speichern")
			self.main_screen_save_as_py_button.configure(text="als Python speichern")
			self.main_screen_save_as_java_button.configure(text="als Java speichern")
			self.main_screen_save_as_kotlin_button.configure(text="als Kotlin speichern")

			self.main_screen_clear_button.configure(text="Texte löschen")
			self.main_screen_edit_text_button.configure(text="Text bearbeiten")
			
			self.main_screen_open_file_button.configure(text="Datei öffnen")
			self.main_screen_open_file_txt_button.configure(text="Textdatei öffnen")
			self.main_screen_open_file_docx_button.configure(text="Word-Datei öffnen")
			self.main_screen_open_file_py_button.configure(text="Python-Datei öffnen")
			
			self.main_screen_coding_button.configure(text="Codierung")
			self.main_screen_coding_function_button.configure(text="kodieren")

			self.main_screen_settings_button.configure(text="Einstellungen")
			self.main_screen_settings_text.configure(text="Einstellungen")
			self.main_screen_settings_language_text.configure(text="Sprachen")

		elif self.main_screen_settings_language_option_varuable == "Español":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_varuable, self.data)

			self.main_screen_save_button.configure(text="guardar texto")
			self.main_screen_save_as_txt_button.configure(text="guardar como txt")
			self.main_screen_save_as_docx_button.configure(text="guardar como docx")
			self.main_screen_save_as_py_button.configure(text="guardar como python")
			self.main_screen_save_as_java_button.configure(text="guardar como java")
			self.main_screen_save_as_kotlin_button.configure(text="guardar como kotlin")

			self.main_screen_clear_button.configure(text="borrar texto")
			self.main_screen_edit_text_button.configure(text="editar texto")
			
			self.main_screen_open_file_button.configure(text="abrir documento")
			self.main_screen_open_file_txt_button.configure(text="abrir archivo de texto")
			self.main_screen_open_file_docx_button.configure(text="abrir archivo de palabra")
			self.main_screen_open_file_py_button.configure(text="abrir archivo python")
			
			self.main_screen_coding_button.configure(text="codificacióng")
			self.main_screen_coding_function_button.configure(text="codificar")

			self.main_screen_settings_button.configure(text="ajustes")
			self.main_screen_settings_text.configure(text="ajustes")
			self.main_screen_settings_language_text.configure(text="lenguas")

	def __autosave_settings__(self):
		self.main_screen_settings_autosave_switch_value = self.main_screen_settings_autosave_value.get()
		if self.main_screen_settings_autosave_switch_value == "on":
			with open("my_diary_autosave_settings.pickle", "wb") as self.autosave_data:
				dump(self.main_screen_settings_autosave_switch_value, self.autosave_data)

			self.main_screen.bind('<Return>', self.__text_autosave__)
			self.main_screen.bind('<Control_L>' + '<s>', self.__text_autosave__)
		
		elif self.main_screen_settings_autosave_switch_value == "off":
			with open("my_diary_autosave_settings.pickle", "wb") as self.autosave_data:
				dump(self.main_screen_settings_autosave_switch_value, self.autosave_data)

			self.main_screen.unbind('<Return>')
			self.main_screen.unbind('<Control_L>' + '<s>')

	def __text_autosave__(self, event):
		self.main_screen_frame_textbox_text_data = self.main_screen_frame_textbox.get("1.0", END)

		with open('my_diary_saved_text.pickle', 'wb') as self.text_data:
			dump(self.main_screen_frame_textbox_text_data, self.text_data)

	def __programm_version__(self, event):
		self.main_screen_additional_window = CTkToplevel()
		self.main_screen_additional_window.geometry("350x105")
		self.main_screen_additional_window.attributes("-toolwindow", True)
		self.main_screen_additional_window.resizable(False, False)
		self.main_screen_additional_window.title("Version_info")

		self.main_screen_additional_window_text = CTkLabel(master=self.main_screen_additional_window, text="версија: 0.5", font=("Roboto Bold", 36))
		self.main_screen_additional_window_text.place(x=0, y=5)

		self.main_screen_additional_window_whats_new_text = CTkLabel(master=self.main_screen_additional_window, text="шта је ново:", height=1)
		self.main_screen_additional_window_whats_new_text.place(x=0, y=47)

		self.main_screen_additional_window_whats_new_1_text = CTkLabel(master=self.main_screen_additional_window, text="1) додат шпански и немачки језци", height=1)
		self.main_screen_additional_window_whats_new_1_text.place(x=75, y=47)

		self.main_screen_additional_window_whats_new_2_text = CTkLabel(master=self.main_screen_additional_window, text="2) Додат аutosave", height=1)
		self.main_screen_additional_window_whats_new_2_text.place(x=75, y=67)

		self.main_screen_additional_window_whats_new_3_text = CTkLabel(master=self.main_screen_additional_window, text="3) Побољшана подршка и модификација за Word", height=1)
		self.main_screen_additional_window_whats_new_3_text.place(x=75, y=87)

	def __theme__(self, event):
		self.main_screen_theme_numbers += 1
		if self.main_screen_theme_numbers % 2 == 0:
			set_appearance_mode("Light")
			self.main_screen.iconbitmap("my_diary_icon2.ico")
		
		else:
			set_appearance_mode("Dark")
			self.main_screen.iconbitmap("my_diary_icon.ico")
	
	def __fullscreen__(self, event):
		self.main_screen_fullscreen_numbers += 1
		if self.main_screen_fullscreen_numbers % 2 == 0:
			self.main_screen.attributes('-fullscreen', True)
			
			self.main_screen_frame.configure(height=842)
			self.main_screen_frame_textbox.configure(height=836)
		
		else:
			self.main_screen.attributes('-fullscreen', False)

			self.main_screen_frame.configure(height=776)
			self.main_screen_frame_textbox.configure(height=770)

	def __exit__(self, event):
		self.main_screen_exit_window = CTkToplevel()
		self.main_screen_exit_window.geometry("350x100")
		self.main_screen_exit_window.attributes("-toolwindow", True)
		self.main_screen_exit_window.resizable(False, False)
		self.main_screen_exit_window.title("exit")

		self.main_screen_exit_window_text = CTkLabel(master=self.main_screen_exit_window, text="желите да изађете?", font=("Roboto Bold", 36))
		self.main_screen_exit_window_text.place(x=0, y=15)

		self.main_screen_exit_yes_button = CTkButton(master=self.main_screen_exit_window, text="да", command=self.__yes_exit__)
		self.main_screen_exit_yes_button.place(x=0, y=60)

		self.main_screen_exit_no_button = CTkButton(master=self.main_screen_exit_window, text="не", command=self.__no_exit__)
		self.main_screen_exit_no_button.place(x=210, y=60)    

	def __yes_exit__(self):
		sys.exit()

	def __no_exit__(self):
		self.main_screen_exit_window.quit()

if __name__ == "__main__":
	programm = Programm() 
	programm.__run__()