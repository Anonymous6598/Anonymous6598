from customtkinter import * 
from tkinter import * 
from pickle import *
from tkinter.messagebox import *
from aspose.words import *
from tkinter import filedialog
from docx.shared import RGBColor, Pt
import os, sys, docx, subprocess

class Program(CTk):

	TITLE = "My Diary"
	ICON = "my diary icon.ico"
	COLOR_THEME = "dark-blue"
	WIDGET_SCALING = 1.251

	def __init__(self, *args, **kwargs) -> None:
		CTk.__init__(self, *args, **kwargs)

		set_widget_scaling(self.WIDGET_SCALING)
		set_default_color_theme(f"{self.COLOR_THEME}")
		deactivate_automatic_dpi_awareness()

		self.title(f"{self.TITLE}")
		self.iconbitmap(f"{self.ICON}")

		self.bind('<F1>', self.__program_version__)  
		self.bind('<Control_L>' + '<f>', self.__fullscreen__) 
		self.bind('<Alt_L>' + '<F4>', self.__exit__)
		self.bind('<Alt_L>' + '<t>', self.__open_terminal__)
		self.bind('<Escape>', self.__close_taskbar_with_peripherals__)
		self.bind('<Control_L>' + '<t>', self.__open_taskbar_with_keybord__)
		self.bind('<Button-3>', self.__right_click_menu__)

		self.protocol("WM_DELETE_WINDOW", self.__exit__)
		
		self.main_screen_fullscreen_numbers = 1

		with open("my_diary_settings.pickle", "rb") as self.data:
			self.new_data = load(self.data)

		with open('my_diary_saved_text.pickle', "rb") as self.text_data:
			self.new_text_data = load(self.text_data)

		with open("my_diary_autosave_settings.pickle", "rb") as self.autosave_data:
			self.new_autosaved_data = load(self.autosave_data)

		with open("my_diary_theme_settings.pickle", "rb") as self.theme_data:
			self.new_theme_data = load(self.theme_data)

		with open("my_diary_text_color.pickle", "rb") as self.text_color_data:
			self.new_text_color_data = load(self.text_color_data)

		with open("my_diary_text_field_color.pickle", "rb") as self.text_field_color_data:
			self.new_text_field_color_data = load(self.text_field_color_data)

		with open("my_diary_text_field_text_height.pickle", "rb") as self.text_field_text_height_data:
			self.new_text_field_text_height_data = load(self.text_field_text_height_data)

		with open("my_diary_button_color.pickle", "rb") as self.button_color_data:
			self.new_button_color_data = load(self.button_color_data)

		self.main_screen_taskbar_button = CTkButton(master=self, text="☰", height=20, width=10, corner_radius=0, fg_color="green", command=self.__open_taskbar__)
		self.main_screen_taskbar_button.grid(column=0, row=0)

		self.main_screen_undo_button = CTkButton(master=self, text="⟲", height=20, width=10, corner_radius=0, fg_color="green", command=self.__undo__)
		self.main_screen_undo_button.grid(column=1, row=0)

		self.main_screen_redo_button = CTkButton(master=self, text="⟳", height=20, width=10, corner_radius=0, fg_color="green", command=self.__redo__)
		self.main_screen_redo_button.grid(column=2, row=0)

		self.main_screen_frame = CTkFrame(master=self, height=769, width=1535, border_width=1, border_color=("black", "white"), corner_radius=0)
		self.main_screen_frame.place(x=0, y=22.35)

		self.bind("<Configure>", self.__frame_resize__)

		self.main_screen_frame_texbox_font = CTkFont(family="Ubuntu", size=22, weight="normal", slant="roman", underline=False, overstrike=False)

		self.main_screen_frame_textbox = CTkTextbox(master=self.main_screen_frame, height=767.5, width=1533.57, corner_radius=0, undo=True, fg_color="transparent", font=self.main_screen_frame_texbox_font, text_color=("black", "white"))
		self.main_screen_frame_textbox.place(x=1, y=1)

		self.main_screen_frame_textbox.bind('<KeyRelease>', self.__word_count__)
		self.main_screen_frame_textbox.bind('<F2>', self.__html_script__)
		self.main_screen_frame_textbox.bind('<KeyPress>', self.__close_taskbar_with_peripherals__)
		self.main_screen_frame_textbox.bind('<ButtonRelease>', self.__close_taskbar_with_peripherals__)

		self.main_screen_taskbar = Program_taskbar(master=self)

		self.main_screen_taskbar_exit_button = CTkButton(master=self.main_screen_taskbar, text="↵", height=20, width=10, corner_radius=5, fg_color="green", command=self.__close_taskbar__)
		self.main_screen_taskbar_exit_button.place(x=278, y=2)

		self.main_screen_save_button = CTkButton(master=self.main_screen_taskbar, text="сачувај текст", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__save__)
		self.main_screen_save_button.place(x=2, y=32)

		self.main_screen_save_as_button = CTkButton(master=self, text="сачувај као ...", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as__)

		self.main_screen_save_as_docx_button = CTkButton(master=self, text="сачувај као docx", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_docx__)

		self.main_screen_clear_button = CTkButton(master=self.main_screen_taskbar, text="обриши текст", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__clear_text__)
		self.main_screen_clear_button.place(x=2, y=67)

		self.main_screen_edit_text_button = CTkButton(master=self.main_screen_taskbar, text="уреди текст", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__edit_text_)
		self.main_screen_edit_text_button.place(x=2, y=102)

		self.main_screen_edit_font_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["System", "Terminal", "Ubuntu", "Script", "Roman", "Modern", "MS Serif"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_size_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["1", "2", "4", "5", "6", "8", "11", "12", "13", "14", "16", "18", "20", "22", "24", "26", "28", "30", "32", "34", "36", "38", "40", "42", "44", "46", "48", "50", "60", "70", "80", "90", "100"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_color_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["black", "white", "red", "green", "blue"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_slant_button =  CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["roman", "italic"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_weight_button =  CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["normal", "bold"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_underline_button =  CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["not underlined", "underlined"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_overstrike_button =  CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["not overstriked", "overstriked"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_open_button = CTkButton(master=self.main_screen_taskbar, text="отвори фајл", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__open__)
		self.main_screen_open_button.place(x=2, y=137)
	
		self.main_screen_open_file_button = CTkButton(master=self, text="отвори ... фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file__)
	
		self.main_screen_open_file_docx_button = CTkButton(master=self, text="отвори docx фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_docx__)
	
		self.main_screen_converter_button = CTkButton(master=self.main_screen_taskbar, text="претварач", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__convert__)
		self.main_screen_converter_button.place(x=2, y=172)

		self.main_screen_pdf_to_word_converter_button = CTkButton(master=self, text="из Pdf у docx", height=20, width=25, corner_radius=0, fg_color="green", command=self.__pdf_to_docx__)

		self.main_screen_word_to_pdf_converter_button = CTkButton(master=self, text="из docx у Pdf", height=20, width=25, corner_radius=0, fg_color="green", command=self.__docx_to_pdf__)

		self.main_screen_pdf_to_txt_converter_button = CTkButton(master=self, text="из Pdf у txt", height=20, width=25, corner_radius=0, fg_color="green", command=self.__pdf_to_txt__)

		self.main_screen_txt_to_pdf_converter_button = CTkButton(master=self, text="из txt у Pdf", height=20, width=25, corner_radius=0, fg_color="green", command=self.__txt_to_pdf__)

		self.main_screen_word_to_txt_converter_button = CTkButton(master=self, text="из docx y txt", height=20, width=25, corner_radius=0, fg_color="green", command=self.__docx_to_txt__)

		self.main_screen_txt_to_word_converter_button = CTkButton(master=self, text="из txt у docx", height=20, width=25, corner_radius=0, fg_color="green", command=self.__txt_to_docx__)

		self.main_screen_code_editor_button = CTkButton(master=self.main_screen_taskbar, text="редактор кода", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__code_editor__)
		self.main_screen_code_editor_button.place(x=2, y=207)

		self.main_screen_open_powershell_button = CTkButton(master=self, text="отвори powershell", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_powershell__)

		self.main_screen_open_terminal_button = CTkButton(master=self, text="отвори терминал", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_shell__)

		self.main_screen_save_code_as_button = CTkButton(master=self, text="сачувај код", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_code_as__)

		self.main_screen_open_code_button = CTkButton(master=self, text="отвори код", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_code__)

		self.main_screen_right_click_menu = Menu(self, tearoff=0)

		self.main_screen_right_click_save_menu = Menu(self, tearoff=0)

		self.main_screen_right_click_open_menu = Menu(self, tearoff=0)

		self.main_screen_right_click_convert_menu = Menu(self, tearoff=0)

		self.main_screen_right_click_code_menu = Menu(self, tearoff=0)

		self.main_screen_settings_button = CTkButton(master=self.main_screen_taskbar, text="подешавања", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__settings__)
		self.main_screen_settings_button.place(x=2, y=242)

		self.main_screen_settings_text = CTkLabel(master=self, text="подешавања", font=("Roboto Bold", 72))

		self.main_screen_settings_language_text = CTkLabel(master=self, text="језици", font=("Roboto Bold", 36))

		self.main_screen_settings_language_option = CTkSegmentedButton(master=self, values=["Српски", "English", "Русский"], selected_color="green", command=self.__language_settings__)

		self.language_data = str(self.new_data)

		self.main_screen_settings_language_option.set(self.language_data)

		self.main_screen_settings_autosave_text = CTkLabel(master=self, text="ауточување", font=("Roboto Bold", 36))

		self.main_screen_settings_autosave_value = StringVar(value=self.new_autosaved_data)

		self.main_screen_settings_autosave_switch_variable = str(self.new_autosaved_data)

		self.main_screen_settings_autosave_switch = CTkSwitch(master=self, text="ауточување", variable=self.main_screen_settings_autosave_value, onvalue="on", offvalue="off", progress_color="green", command=self.__autosave_settings__)

		self.autosaved_text = str(self.new_text_data)

		self.main_screen_settings_autosave_switch_value = self.main_screen_settings_autosave_value.get()
			
		self.main_screen_settings_theme_mode_text = CTkLabel(master=self, text="тема", font=("Roboto Bold", 36))

		self.main_screen_settings_theme_mode_option = CTkSegmentedButton(master=self, values=["system", "dark", "light"], selected_color="green", command=self.__theme_settings__)

		self.theme = str(self.new_theme_data)

		self.main_screen_settings_theme_mode_option.set(self.theme)

		self.main_screen_settings_customatization_text = CTkLabel(master=self, text="Спољни изглед", font=("Roboto Bold", 36))

		self.main_screen_settings_customatization_table = CTkTabview(master=self, height=50, width=400, border_width=1, border_color=("black", "white"), segmented_button_selected_color="green", text_color="white")

		self.main_screen_settings_customatization_table.add("Text")
		self.main_screen_settings_customatization_table.add("Text field")
		self.main_screen_settings_customatization_table.add("Button")

		self.main_screen_settings_customatization_text_color_text = CTkLabel(master=self.main_screen_settings_customatization_table.tab("Text"), text="Боја текста", font=("Roboto Bold", 36))
		self.main_screen_settings_customatization_text_color_text.grid(column=0, row=0)

		self.main_screen_settings_customatization_text_color_option = CTkSegmentedButton(master=self.main_screen_settings_customatization_table.tab("Text"), values=["red", "blue", "green", "black", "white"], selected_color="green", command=self.__change_text_color__)
		self.main_screen_settings_customatization_text_color_option.grid(column=0, row=1)


		self.main_screen_settings_customatization_text_field_color_text = CTkLabel(master=self.main_screen_settings_customatization_table.tab("Text field"), text="Боја поља", font=("Roboto Bold", 36))
		self.main_screen_settings_customatization_text_field_color_text.grid(column=0, row=0)

		self.main_screen_settings_customatization_text_field_color_option = CTkSegmentedButton(master=self.main_screen_settings_customatization_table.tab("Text field"), values=["red", "blue", "green", "black", "white", "transparent"], selected_color="green", command=self.__change_text_field_color__)
		self.main_screen_settings_customatization_text_field_color_option.grid(column=0, row=1)

		self.main_screen_settings_customatization_text_field_height_text = CTkLabel(master=self.main_screen_settings_customatization_table.tab("Text field"), text="Висина текста поља", font=("Roboto Bold", 36))
		self.main_screen_settings_customatization_text_field_height_text.grid(column=0, row=2)

		self.main_screen_settings_customatization_text_field_height_option = CTkSegmentedButton(master=self.main_screen_settings_customatization_table.tab("Text field"), values=["14", "22", "28", "32"], selected_color="green", command=self.__change_text_field_text_height__)
		self.main_screen_settings_customatization_text_field_height_option.grid(column=0, row=3)


		self.main_screen_settings_customatization_button_color_text = CTkLabel(master=self.main_screen_settings_customatization_table.tab("Button"), text="Боја дугма", font=("Roboto Bold", 36))
		self.main_screen_settings_customatization_button_color_text.grid(column=0, row=0)

		self.main_screen_settings_customatization_button_color_option = CTkSegmentedButton(master=self.main_screen_settings_customatization_table.tab("Button"), values=["red", "blue", "green", "black", "orange", "yellow", "purple"], selected_color="green", command=self.__change_button_color__)
		self.main_screen_settings_customatization_button_color_option.grid(column=0, row=1)

		self.text_color = str(self.new_text_color_data)

		self.text_field_color = str(self.new_text_field_color_data)

		self.text_field_text_height = int(self.new_text_field_text_height_data)

		self.button_color = str(self.new_button_color_data)

		self.main_screen_settings_customatization_text_color_option.set(self.text_color)

		self.main_screen_settings_customatization_text_field_color_option.set(self.text_field_color)

		self.main_screen_settings_customatization_text_field_height_option.set(str(self.text_field_text_height))

		self.main_screen_settings_customatization_button_color_option.set(self.button_color)

		if self.language_data == "Српски":
			self.main_screen_save_button.configure(text="сачувај текст")
			self.main_screen_save_as_button.configure(text="сачувај као ...")
			self.main_screen_save_as_docx_button.configure(text="сачувај као docx")
	
			self.main_screen_clear_button.configure(text="обриши текст")
			self.main_screen_edit_text_button.configure(text="уреди текст")
			
			self.main_screen_open_button.configure(text="отвори фајл")
			self.main_screen_open_file_button.configure(text="отвори ... фајл")
			self.main_screen_open_file_docx_button.configure(text="отвори docx фајл")

			self.main_screen_converter_button.configure(text="претварач")
			self.main_screen_pdf_to_word_converter_button.configure(text="из Pdf у docx")
			self.main_screen_word_to_pdf_converter_button.configure(text="из docx у Pdf")
			self.main_screen_pdf_to_txt_converter_button.configure(text="из Pdf у txt")
			self.main_screen_txt_to_pdf_converter_button.configure(text="из txt у Pdf")
			self.main_screen_word_to_txt_converter_button.configure(text="из docx у txt")
			self.main_screen_txt_to_word_converter_button.configure(text="из txt у docx")

			self.main_screen_settings_button.configure(text="подешавања")
			self.main_screen_settings_text.configure(text="подешавања")
			self.main_screen_settings_language_text.configure(text="језици")
			self.main_screen_settings_autosave_text.configure(text="ауточување")
			self.main_screen_settings_autosave_switch.configure(text="ауточување")
			self.main_screen_settings_theme_mode_text.configure(text="теме")
			self.main_screen_settings_customatization_text.configure(text="Спољни изглед")
			self.main_screen_settings_customatization_text_color_text.configure(text="Боја текста")
			self.main_screen_settings_customatization_text_field_color_text.configure(text="Боја поља")
			self.main_screen_settings_customatization_text_field_height_text.configure(text="Висина текста поља")
			self.main_screen_settings_customatization_button_color_text.configure(text="Боја дугма")

			self.main_screen_code_editor_button.configure(text="кодови редактор")
			self.main_screen_open_powershell_button.configure(text="отвори powershell")
			self.main_screen_open_terminal_button.configure(text="отвори терминал")
			self.main_screen_save_code_as_button.configure(text="сачувај код као")
			self.main_screen_open_code_button.configure(text="отвори код")

			self.main_screen_right_click_menu.add_cascade(label="сачувај као", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_save_menu)
			self.main_screen_right_click_save_menu.add_command(label="...", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as__)
			self.main_screen_right_click_save_menu.add_command(label=".docx", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_docx__)

			self.main_screen_right_click_menu.add_cascade(label="отвори као", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_open_menu)
			self.main_screen_right_click_open_menu.add_command(label="...", font=("Roman", 16), background="green", foreground="white", command=self.__open_file__)
			self.main_screen_right_click_open_menu.add_command(label=".docx", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_docx__)

			self.main_screen_right_click_menu.add_cascade(label="конвертирај", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_convert_menu)
			self.main_screen_right_click_convert_menu.add_command(label="из docx у pdf", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="из pdf у docx", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="из txt у pdf", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="из pdf у txt", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_txt__)
			self.main_screen_right_click_convert_menu.add_command(label="из txt у docx", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="из docx у txt", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_txt__)

			self.main_screen_right_click_menu.add_cascade(label="кодирај", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_code_menu)
			self.main_screen_right_click_code_menu.add_command(label="отвори powershell", font=("Roman", 16), background="green", foreground="white", command=self.__open_powershell__)
			self.main_screen_right_click_code_menu.add_command(label="отвори терминал", font=("Roman", 16), background="green", foreground="white", command=self.__open_shell__)
			self.main_screen_right_click_code_menu.add_command(label="сачувај код као", font=("Roman", 16), background="green", foreground="white", command=self.__save_code_as__)
			self.main_screen_right_click_code_menu.add_command(label="отвори код", font=("Roman", 16), background="green", foreground="white", command=self.__open_code__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="обриши текст", font=("Roman", 16), background="green", foreground="white", command=self.__clear_text__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="копирај", font=("Roman", 16), background="green", foreground="white", command=self.__copy__)
			self.main_screen_right_click_menu.add_command(label="залепи", font=("Roman", 16), background="green", foreground="white", command=self.__paste__)
			self.main_screen_right_click_menu.add_command(label="исеци", font=("Roman", 16), background="green", foreground="white", command=self.__cut__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="изаћи", font=("Roman", 16), background="green", foreground="white", command=self.__exit__)

		elif self.language_data == "English":
			self.main_screen_save_button.configure(text="save text")
			self.main_screen_save_as_button.configure(text="save as ...")
			self.main_screen_save_as_docx_button.configure(text="save as docx")

			self.main_screen_clear_button.configure(text="clear text")
			self.main_screen_edit_text_button.configure(text="edit text")
			
			self.main_screen_open_button.configure(text="open file")
			self.main_screen_open_file_button.configure(text="open ... file")
			self.main_screen_open_file_docx_button.configure(text="open docx file")

			self.main_screen_converter_button.configure(text="converterer")
			self.main_screen_pdf_to_word_converter_button.configure(text="from Pdf to docx")
			self.main_screen_word_to_pdf_converter_button.configure(text="from docx to Pdf")
			self.main_screen_pdf_to_txt_converter_button.configure(text="from Pdf to txt")
			self.main_screen_txt_to_pdf_converter_button.configure(text="from txt to Pdf")
			self.main_screen_word_to_txt_converter_button.configure(text="from docx to txt")
			self.main_screen_txt_to_word_converter_button.configure(text="from txt to docx")

			self.main_screen_code_editor_button.configure(text="code editor")
			self.main_screen_open_powershell_button.configure(text="open powershell")
			self.main_screen_open_terminal_button.configure(text="open terminal")
			self.main_screen_save_code_as_button.configure(text="save code as")
			self.main_screen_open_code_button.configure(text="open code")

			self.main_screen_settings_button.configure(text="settings")
			self.main_screen_settings_text.configure(text="settings")
			self.main_screen_settings_language_text.configure(text="languages")
			self.main_screen_settings_autosave_text.configure(text="autosave")
			self.main_screen_settings_autosave_switch.configure(text="autosave")
			self.main_screen_settings_theme_mode_text.configure(text="theme")
			self.main_screen_settings_customatization_text.configure(text="Customatization")
			self.main_screen_settings_customatization_text_color_text.configure(text="Text color")
			self.main_screen_settings_customatization_text_field_color_text.configure(text="Text field color")
			self.main_screen_settings_customatization_text_field_height_text.configure(text="Text field text height")
			self.main_screen_settings_customatization_button_color_text.configure(text="Button color")

			self.main_screen_right_click_menu.add_cascade(label="save as", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_save_menu)
			self.main_screen_right_click_save_menu.add_command(label="...", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as__)
			self.main_screen_right_click_save_menu.add_command(label=".docx", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_docx__)

			self.main_screen_right_click_menu.add_cascade(label="open as", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_open_menu)
			self.main_screen_right_click_open_menu.add_command(label="...", font=("Roman", 16), background="green", foreground="white", command=self.__open_file__)
			self.main_screen_right_click_open_menu.add_command(label=".docx", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_docx__)

			self.main_screen_right_click_menu.add_cascade(label="convert", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_convert_menu)
			self.main_screen_right_click_convert_menu.add_command(label="from docx to pdf", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="from pdf to docx", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="from txt to pdf", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="from pdf to txt", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_txt__)
			self.main_screen_right_click_convert_menu.add_command(label="from txt to docx", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="from docx to txt", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_txt__)

			self.main_screen_right_click_menu.add_cascade(label="code", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_code_menu)
			self.main_screen_right_click_code_menu.add_command(label="open powershell", font=("Roman", 16), background="green", foreground="white", command=self.__open_powershell__)
			self.main_screen_right_click_code_menu.add_command(label="open terminal", font=("Roman", 16), background="green", foreground="white", command=self.__open_shell__)
			self.main_screen_right_click_code_menu.add_command(label="save code as", font=("Roman", 16), background="green", foreground="white", command=self.__save_code_as__)
			self.main_screen_right_click_code_menu.add_command(label="open code", font=("Roman", 16), background="green", foreground="white", command=self.__open_code__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="delete text", font=("Roman", 16), background="green", foreground="white", command=self.__clear_text__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="copy", font=("Roman", 16), background="green", foreground="white", command=self.__copy__)
			self.main_screen_right_click_menu.add_command(label="paste", font=("Roman", 16), background="green", foreground="white", command=self.__paste__)
			self.main_screen_right_click_menu.add_command(label="cut", font=("Roman", 16), background="green", foreground="white", command=self.__cut__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="exit", font=("Roman", 16), background="green", foreground="white", command=self.__exit__)

		else:
			self.main_screen_save_button.configure(text="сохранить текст")
			self.main_screen_save_as_button.configure(text="сохранить как ...")
			self.main_screen_save_as_docx_button.configure(text="сохранить как ворд")

			self.main_screen_clear_button.configure(text="удалить текст")
			self.main_screen_edit_text_button.configure(text="редактировать текст")
			
			self.main_screen_open_button.configure(text="открыть файл")
			self.main_screen_open_file_button.configure(text="открыть ... файл")
			self.main_screen_open_file_docx_button.configure(text="открыть ворд файл")

			self.main_screen_converter_button.configure(text="конвертер")
			self.main_screen_pdf_to_word_converter_button.configure(text="из Пдф в ворд")
			self.main_screen_word_to_pdf_converter_button.configure(text="из ворд в Пдф")
			self.main_screen_pdf_to_txt_converter_button.configure(text="из Пдф в текстовый файл")
			self.main_screen_txt_to_pdf_converter_button.configure(text="из текстового файла в Пдф")
			self.main_screen_word_to_txt_converter_button.configure(text="из ворд в текстовый файл")
			self.main_screen_txt_to_word_converter_button.configure(text="из текстового файла в ворд")

			self.main_screen_code_editor_button.configure(text="кодовый редактор")
			self.main_screen_open_powershell_button.configure(text="открыть powershell")
			self.main_screen_open_terminal_button.configure(text="открыть терминал")
			self.main_screen_save_code_as_button.configure(text="сохранить код как")
			self.main_screen_open_code_button.configure(text="открыть код")

			self.main_screen_settings_button.configure(text="настройки")
			self.main_screen_settings_text.configure(text="настройки")
			self.main_screen_settings_language_text.configure(text="языки")
			self.main_screen_settings_autosave_text.configure(text="автосохранение")
			self.main_screen_settings_autosave_switch.configure(text="автосохранение")
			self.main_screen_settings_theme_mode_text.configure(text="темы")
			self.main_screen_settings_customatization_text.configure(text="Внешний вид")
			self.main_screen_settings_customatization_text_color_text.configure(text="Цвет текста")
			self.main_screen_settings_customatization_text_field_color_text.configure(text="Цвет поля")
			self.main_screen_settings_customatization_text_field_height_text.configure(text="Высота текста поля")
			self.main_screen_settings_customatization_button_color_text.configure(text="Цвет кнопки")

			self.main_screen_right_click_menu.add_cascade(label="сохранить как", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_save_menu)
			self.main_screen_right_click_save_menu.add_command(label="...", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as__)
			self.main_screen_right_click_save_menu.add_command(label="ворд", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_docx__)
			
			self.main_screen_right_click_menu.add_cascade(label="открыть как", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_open_menu)
			self.main_screen_right_click_open_menu.add_command(label="...", font=("Roman", 16), background="green", foreground="white", command=self.__open_file__)
			self.main_screen_right_click_open_menu.add_command(label="ворд", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_docx__)

			self.main_screen_right_click_menu.add_cascade(label="конвертировать", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_convert_menu)
			self.main_screen_right_click_convert_menu.add_command(label="из ворд в Пдф", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="из Пдф в ворд", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="из текста в Пдф", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="из Пдф в текст", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_txt__)
			self.main_screen_right_click_convert_menu.add_command(label="из текста в ворд", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="из ворда в текст", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_txt__)

			self.main_screen_right_click_menu.add_cascade(label="кодирование", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_code_menu)
			self.main_screen_right_click_code_menu.add_command(label="открыть терминал", font=("Roman", 16), background="green", foreground="white", command=self.__open_powershell__)
			self.main_screen_right_click_code_menu.add_command(label="открыть терминал", font=("Roman", 16), background="green", foreground="white", command=self.__open_shell__)
			self.main_screen_right_click_code_menu.add_command(label="сохранить код как", font=("Roman", 16), background="green", foreground="white", command=self.__save_code_as__)
			self.main_screen_right_click_code_menu.add_command(label="открыть код", font=("Roman", 16), background="green", foreground="white", command=self.__open_code__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="удалить текст", font=("Roman", 16), background="green", foreground="white", command=self.__clear_text__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="копировать", font=("Roman", 16), background="green", foreground="white", command=self.__copy__)
			self.main_screen_right_click_menu.add_command(label="вставить", font=("Roman", 16), background="green", foreground="white", command=self.__paste__)
			self.main_screen_right_click_menu.add_command(label="вырезать", font=("Roman", 16), background="green", foreground="white", command=self.__cut__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="выход", font=("Roman", 16), background="green", foreground="white", command=self.__exit__)


		if self.main_screen_settings_autosave_switch_variable == "on":
			self.main_screen_frame_textbox.insert("1.0", self.autosaved_text)

			self.bind('<Return>', self.__text_autosave__)
			self.bind('<Control_L>' + '<s>', self.__text_autosave__)

		else:
			self.unbind('<Return>')
			self.unbind('<Control_L>' + '<s>')


		set_appearance_mode(self.theme)


		self.main_screen_taskbar_button.configure(text_color=self.text_color)
		self.main_screen_undo_button.configure(text_color=self.text_color)
		self.main_screen_redo_button.configure(text_color=self.text_color)
		self.main_screen_taskbar_exit_button.configure(text_color=self.text_color)

		self.main_screen_save_button.configure(text_color=self.text_color)
		self.main_screen_save_as_button.configure(text_color=self.text_color)
		self.main_screen_save_as_docx_button.configure(text_color=self.text_color)

		self.main_screen_clear_button.configure(text_color=self.text_color)
		self.main_screen_edit_text_button.configure(text_color=self.text_color)
		self.main_screen_edit_font_button.configure(text_color=self.text_color, dropdown_text_color=self.text_color)
		self.main_screen_edit_size_button.configure(text_color=self.text_color, dropdown_text_color=self.text_color)
		self.main_screen_edit_color_button.configure(text_color=self.text_color, dropdown_text_color=self.text_color)
		self.main_screen_edit_slant_button.configure(text_color=self.text_color, dropdown_text_color=self.text_color)
		self.main_screen_edit_weight_button.configure(text_color=self.text_color, dropdown_text_color=self.text_color)
		self.main_screen_edit_underline_button.configure(text_color=self.text_color, dropdown_text_color=self.text_color)
		self.main_screen_edit_overstrike_button.configure(text_color=self.text_color, dropdown_text_color=self.text_color)

		self.main_screen_open_button.configure(text_color=self.text_color)
		self.main_screen_open_file_button.configure(text_color=self.text_color)
		self.main_screen_open_file_docx_button.configure(text_color=self.text_color)

		self.main_screen_converter_button.configure(text_color=self.text_color)
		self.main_screen_pdf_to_word_converter_button.configure(text_color=self.text_color)
		self.main_screen_word_to_pdf_converter_button.configure(text_color=self.text_color)
		self.main_screen_pdf_to_txt_converter_button.configure(text_color=self.text_color)
		self.main_screen_txt_to_pdf_converter_button.configure(text_color=self.text_color)
		self.main_screen_word_to_txt_converter_button.configure(text_color=self.text_color)
		self.main_screen_txt_to_word_converter_button.configure(text_color=self.text_color)

		self.main_screen_code_editor_button.configure(text_color=self.text_color)
		self.main_screen_open_powershell_button.configure(text_color=self.text_color)
		self.main_screen_save_code_as_button.configure(text_color=self.text_color)
		self.main_screen_open_code_button.configure(text_color=self.text_color)

		self.main_screen_settings_button.configure(text_color=self.text_color)
		self.main_screen_settings_text.configure(text_color=self.text_color)
		self.main_screen_settings_language_text.configure(text_color=self.text_color)
		self.main_screen_settings_autosave_text.configure(text_color=self.text_color)
		self.main_screen_settings_autosave_switch.configure(text_color=self.text_color)
		self.main_screen_settings_theme_mode_text.configure(text_color=self.text_color)
		self.main_screen_settings_customatization_text.configure(text_color=self.text_color)
		self.main_screen_settings_customatization_text_color_text.configure(text_color=self.text_color)
		self.main_screen_settings_customatization_text_field_color_text.configure(text_color=self.text_color)
		self.main_screen_settings_customatization_text_field_height_text.configure(text_color=self.text_color)
		self.main_screen_settings_customatization_button_color_text.configure(text_color=self.text_color)

		
		self.main_screen_frame_textbox.configure(fg_color=self.text_field_color)


		self.main_screen_frame_texbox_font.configure(size=self.text_field_text_height)


		self.main_screen_taskbar_button.configure(fg_color=self.button_color)
		self.main_screen_undo_button.configure(fg_color=self.button_color)
		self.main_screen_redo_button.configure(fg_color=self.button_color)
		self.main_screen_taskbar_exit_button.configure(fg_color=self.button_color)

		self.main_screen_save_button.configure(fg_color=self.button_color)
		self.main_screen_save_as_button.configure(fg_color=self.button_color)
		self.main_screen_save_as_docx_button.configure(fg_color=self.button_color)

		self.main_screen_clear_button.configure(fg_color=self.button_color)
		self.main_screen_edit_text_button.configure(fg_color=self.button_color)
		self.main_screen_edit_font_button.configure(fg_color=self.button_color, border_color=self.button_color, button_color=self.button_color, button_hover_color=self.button_color, dropdown_fg_color=self.button_color)
		self.main_screen_edit_size_button.configure(fg_color=self.button_color, border_color=self.button_color, button_color=self.button_color, button_hover_color=self.button_color, dropdown_fg_color=self.button_color)
		self.main_screen_edit_color_button.configure(fg_color=self.button_color, border_color=self.button_color, button_color=self.button_color, button_hover_color=self.button_color, dropdown_fg_color=self.button_color)
		self.main_screen_edit_slant_button.configure(fg_color=self.button_color, border_color=self.button_color, button_color=self.button_color, button_hover_color=self.button_color, dropdown_fg_color=self.button_color)
		self.main_screen_edit_weight_button.configure(fg_color=self.button_color, border_color=self.button_color, button_color=self.button_color, button_hover_color=self.button_color, dropdown_fg_color=self.button_color)
		self.main_screen_edit_underline_button.configure(fg_color=self.button_color, border_color=self.button_color, button_color=self.button_color, button_hover_color=self.button_color, dropdown_fg_color=self.button_color)
		self.main_screen_edit_overstrike_button.configure(fg_color=self.button_color, border_color=self.button_color, button_color=self.button_color, button_hover_color=self.button_color, dropdown_fg_color=self.button_color)

		self.main_screen_open_button.configure(fg_color=self.button_color)
		self.main_screen_open_file_button.configure(fg_color=self.button_color)
		self.main_screen_open_file_docx_button.configure(fg_color=self.button_color)
	
		self.main_screen_converter_button.configure(fg_color=self.button_color)
		self.main_screen_pdf_to_word_converter_button.configure(fg_color=self.button_color)
		self.main_screen_word_to_pdf_converter_button.configure(fg_color=self.button_color)
		self.main_screen_pdf_to_txt_converter_button.configure(fg_color=self.button_color)
		self.main_screen_word_to_txt_converter_button.configure(fg_color=self.button_color)
		self.main_screen_txt_to_word_converter_button.configure(fg_color=self.button_color)

		self.main_screen_code_editor_button.configure(fg_color=self.button_color)
		self.main_screen_open_powershell_button.configure(fg_color=self.button_color)
		self.main_screen_save_code_as_button.configure(fg_color=self.button_color)
		self.main_screen_open_code_button.configure(fg_color=self.button_color)
	
		self.main_screen_settings_button.configure(fg_color=self.button_color)
		self.main_screen_settings_language_option.configure(selected_color=self.button_color)
		self.main_screen_settings_autosave_switch.configure(progress_color=self.button_color)
		self.main_screen_settings_theme_mode_option.configure(selected_color=self.button_color)
		self.main_screen_settings_customatization_table.configure(segmented_button_selected_color=self.button_color)
		self.main_screen_settings_customatization_text_color_option.configure(selected_color=self.button_color)
		self.main_screen_settings_customatization_text_field_color_option.configure(selected_color=self.button_color)
		self.main_screen_settings_customatization_text_field_height_option.configure(selected_color=self.button_color)
		self.main_screen_settings_customatization_button_color_option.configure(selected_color=self.button_color)


		self.main_screen_word_counter_variable = 0

		self.main_screen_word_counter_data_variable = IntVar(value=self.main_screen_word_counter_variable)

		self.main_screen_word_counter = CTkLabel(master=self, width=25, height=20, corner_radius=0, textvariable=self.main_screen_word_counter_data_variable)
		self.main_screen_word_counter.place(x=1500, y=0)
				 
	def __frame_resize__(self, event) -> None:
		if self.winfo_width() <= 958:
			self.main_screen_frame.configure(width=1535 / 2 - 2)
			self.main_screen_frame_textbox.configure(width=1535 / 2 - 5)

		else:
			self.main_screen_frame.configure(width=1535)
			self.main_screen_frame_textbox.configure(width=1533.57)

	def __open_taskbar__(self) -> None:
		self.main_screen_taskbar.grid(column=0, row=0)
		self.main_screen_taskbar_button.grid_forget()

	def __open_taskbar_with_keybord__(self, event) -> None:
		self.main_screen_taskbar.grid(column=0, row=0)
		self.main_screen_taskbar_button.grid_forget()

	def __close_taskbar__(self) -> None:
		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_taskbar.grid_forget()

	def __close_taskbar_with_peripherals__(self, event) -> None:
		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_taskbar.grid_forget()

	def __quit_operation__(self) -> None:
		self.bind('<Control_L>' + '<Tab>' + '<t>', self.__open_taskbar_with_keybord__)

		self.main_screen_taskbar_button.configure(text="☰", command=self.__open_taskbar__)

		self.main_screen_frame.configure(height=769)
		self.main_screen_frame_textbox.configure(height=767.5)

		self.main_screen_undo_button.grid(column=1, row=0)
		self.main_screen_redo_button.grid(column=2, row=0)
		self.main_screen_frame.place(x=0, y=22.35)
		self.main_screen_word_counter.place(x=1500, y=0)

		self.main_screen_save_as_button.grid_forget()
		self.main_screen_save_as_docx_button.grid_forget()

		self.main_screen_edit_font_button.grid_forget()
		self.main_screen_edit_size_button.grid_forget()
		self.main_screen_edit_color_button.grid_forget()
		self.main_screen_edit_slant_button.grid_forget()
		self.main_screen_edit_weight_button.grid_forget()
		self.main_screen_edit_underline_button.grid_forget()
		self.main_screen_edit_overstrike_button.grid_forget()

		self.main_screen_open_file_button.grid_forget()
		self.main_screen_open_file_docx_button.grid_forget()

		self.main_screen_pdf_to_word_converter_button.grid_forget()
		self.main_screen_word_to_pdf_converter_button.grid_forget()
		self.main_screen_pdf_to_txt_converter_button.grid_forget()
		self.main_screen_txt_to_pdf_converter_button.grid_forget()
		self.main_screen_word_to_txt_converter_button.grid_forget()
		self.main_screen_txt_to_word_converter_button.grid_forget()

		self.main_screen_open_powershell_button.grid_forget()
		self.main_screen_open_terminal_button.grid_forget()
		self.main_screen_save_code_as_button.grid_forget()
		self.main_screen_open_code_button.grid_forget()

		self.main_screen_settings_text.place_forget()
		self.main_screen_settings_language_text.place_forget()
		self.main_screen_settings_language_option.place_forget()						          
		self.main_screen_settings_autosave_text.place_forget()
		self.main_screen_settings_autosave_switch.place_forget()
		self.main_screen_settings_theme_mode_text.place_forget()
		self.main_screen_settings_theme_mode_option.place_forget()
		self.main_screen_settings_customatization_text.place_forget()
		self.main_screen_settings_customatization_table.place_forget()

	def __undo__(self) -> None:
		try:
			self.main_screen_frame_textbox.edit_undo()

		except TclError:
			pass

	def __redo__(self) -> None:
		try:
			self.main_screen_frame_textbox.edit_redo()

		except TclError:
			pass

	def __save__(self) -> None:
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_save_as_button.grid(column=1, row=0)
		self.main_screen_save_as_docx_button.grid(column=2, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

		self.main_screen_undo_button.grid_forget()
		self.main_screen_redo_button.grid_forget()
		self.main_screen_word_counter.place_forget()

	def __save_text_as__(self) -> None:
		try:
			with open(filedialog.asksaveasfilename(filetypes=[("All Files (*.*)", "*.*"), ("Text file (*.txt)", "*.txt"), ("Python file (*.py)", "*.py"), ("Java file (*.java)", "*.java"), ("C# file (*.cs)", "*.cs"), ("HTML file (*.html)", "*.html"), ("CSS file (*.css)", "*.css"), ("JavaScript file (*.js)", "*.js"), ("C++ file (*.cpp)", "*.cpp")], defaultextension=[("All Files (*.*)", "*.*"), ("Text file (*.txt)", "*.txt"), ("Python file (*.py)", "*.py"), ("Java file (*.java)", "*.java"), ("C# file (*.cs)", "*.cs"), ("HTML file (*.html)", "*.html"), ("CSS file (*.css)", "*.css"), ("JavaScript file (*.js)", "*.js"), ("C++ file (*.cpp)", "*.cpp")]), "w+", encoding="UTF-8") as self.file:
				self.file_data = self.main_screen_frame_textbox.get("1.0", END)
				self.file.write(self.file_data)
				self.file.close()

		except FileNotFoundError:
			pass
		
	def __save_text_as_docx__(self) -> None:
		try:
			self.file = docx.Document()
			self.file_data = self.main_screen_frame_textbox.get("1.0", END)
			self.file_run = self.file.add_paragraph().add_run(self.file_data)
			self.font = self.file_run.font
			self.font.name = str(self.main_screen_edit_font_button_data)
			self.font.size = Pt(int(self.main_screen_edit_size_button_data))
			if self.main_screen_edit_color_button_data == "black":
				self.font.color.rgb = RGBColor(0, 0, 0)

			elif self.main_screen_edit_color_button_data == "white":
				self.font.color.rgb = RGBColor(255, 255, 255)

			elif self.main_screen_edit_color_button_data == "red":
				self.font.color.rgb = RGBColor(250, 0, 0)

			elif self.main_screen_edit_color_button_data == "green":
				self.font.color.rgb = RGBColor(0, 255, 0)

			else:
				self.font.color.rgb = RGBColor(0, 0, 255)

						
			if self.main_screen_edit_slant_button_data == "italic":
				self.file_run.italic = True

			else:
				self.file_run.italic = False


			if self.main_screen_edit_weight_button_data == "bold":
				self.file_run.bold = True

			else:
				self.file_run.bold = False


			if self.main_screen_edit_underline_button_data == "underlined":
				self.file_run.underline = True

			else:
				self.file_run.underline = False


			if self.main_screen_edit_overstrike_button_data == "overstriked":
				self.font.strike = True

			else:
				self.font.strike = False

			self.file.save(filedialog.asksaveasfilename(filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")]))

		except FileNotFoundError:
			pass

		except AttributeError:
			if self.language_data == "Српски":
				showwarning(title="Пажња", message="Заборавили сте да отредактирате текст")

			elif self.language_data == "English":
				showwarning(title="Warning", message="You forgot to edit your text")

			else:
				showwarning(title="Внимание", message="Вы забыли отредактировать свой текст") 
			
	def __clear_text__(self) -> None:
		self.main_screen_frame_textbox.delete("1.0", END)

	def __edit_text_(self) -> None:
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_edit_font_button.grid(column=1, row=0)
		self.main_screen_edit_size_button.grid(column=2, row=0)
		self.main_screen_edit_color_button.grid(column=3, row=0)
		self.main_screen_edit_slant_button.grid(column=4, row=0)
		self.main_screen_edit_weight_button.grid(column=5, row=0)
		self.main_screen_edit_underline_button.grid(column=6, row=0)
		self.main_screen_edit_overstrike_button.grid(column=7, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()
		
		self.main_screen_undo_button.grid_forget()
		self.main_screen_redo_button.grid_forget()
		self.main_screen_word_counter.place_forget()

	def __edit_text_font__(self, configure) -> None:
		self.main_screen_edit_font_button_data = self.main_screen_edit_font_button.get()															 
		self.main_screen_edit_size_button_data = self.main_screen_edit_size_button.get()
		self.main_screen_edit_color_button_data = self.main_screen_edit_color_button.get()
		self.main_screen_edit_slant_button_data = self.main_screen_edit_slant_button.get()
		self.main_screen_edit_weight_button_data = self.main_screen_edit_weight_button.get()
		self.main_screen_edit_underline_button_data = self.main_screen_edit_underline_button.get()
		self.main_screen_edit_overstrike_button_data = self.main_screen_edit_overstrike_button.get()

		self.main_screen_frame_texbox_font.configure(family=self.main_screen_edit_font_button_data, size=int(self.main_screen_edit_size_button_data), slant=self.main_screen_edit_slant_button_data, weight=self.main_screen_edit_weight_button_data)
		self.main_screen_frame_textbox.configure(text_color=self.main_screen_edit_color_button_data)

		if self.main_screen_edit_underline_button_data == "not underlined":		
			self.main_screen_frame_texbox_font.configure(underline=False)			

		else:		
			self.main_screen_frame_texbox_font.configure(underline=True)
		
		if self.main_screen_edit_overstrike_button_data == "not overstriked":
			self.main_screen_frame_texbox_font.configure(overstrike=False)

		else:		
			self.main_screen_frame_texbox_font.configure(overstrike=True)

	def __open__(self) -> None:
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_open_file_button.grid(column=1, row=0)
		self.main_screen_open_file_docx_button.grid(column=2, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

		self.main_screen_undo_button.grid_forget()
		self.main_screen_redo_button.grid_forget()
		self.main_screen_word_counter.place_forget()

	def __open_file__(self) -> None:
		try:
			with open(filedialog.askopenfilename(title="open file", filetypes=[("All Files (*.*)", "*.*"), ("Text file (*.txt)", "*.txt"), ("Python file (*.py)", "*.py"), ("Java file (*.java)", "*.java"), ("C# file (*.cs)", "*.cs"), ("HTML file (*.html)", "*.html"), ("CSS file (*.css)", "*.css"), ("JavaScript file (*.js)", "*.js"), ("C++ file (*.cpp)", "*.cpp")], defaultextension=[("All Files (*.*)", "*.*"), ("Text file (*.txt)", "*.txt"), ("Python file (*.py)", "*.py"), ("Java file (*.java)", "*.java"), ("C# file (*.cs)", "*.cs"), ("HTML file (*.html)", "*.html"), ("CSS file (*.css)", "*.css"), ("JavaScript file (*.js)", "*.js"), ("C++ file (*.cpp)", "*.cpp")]), "r+", encoding="UTF-8") as self.openned_file:
				self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

		except FileNotFoundError:
			pass

	def __open_file_docx__(self) -> None:
		try:
			self.openned_file = docx.Document(filedialog.askopenfilename(title="open docx file", filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")]))
			self.openned_file_data = []
			for self.paragraphs in self.openned_file.paragraphs:
				self.openned_file_data.append(self.paragraphs.text)

			self.main_screen_frame_textbox.insert("1.0", "\n".join(self.openned_file_data))

		except docx.opc.exceptions.PackageNotFoundError:
			pass

	def __convert__(self) -> None:
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_pdf_to_word_converter_button.grid(column=1, row=0)
		self.main_screen_word_to_pdf_converter_button.grid(column=2, row=0)
		self.main_screen_pdf_to_txt_converter_button.grid(column=3, row=0)
		self.main_screen_txt_to_pdf_converter_button.grid(column=4, row=0)
		self.main_screen_word_to_txt_converter_button.grid(column=5, row=0)
		self.main_screen_txt_to_word_converter_button.grid(column=6, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

		self.main_screen_undo_button.grid_forget()
		self.main_screen_redo_button.grid_forget()
		self.main_screen_word_counter.place_forget()

	def __pdf_to_docx__(self) -> None:
		try:
			self.file = filedialog.askopenfilename(title="convert pdf file", filetypes=[("Pdf file (*.pdf)", "*.pdf")], defaultextension=[("Pdf file (*.pdf)", "*.pdf")])
			self.converted_file = Document(self.file)
			self.converted_file.save(self.file + ".docx")

		except RuntimeError:
			pass
		
	def __docx_to_pdf__(self) -> None:
		try:
			self.file =  filedialog.askopenfilename(title="convert docx file", filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")])
			self.converted_file = Document(self.file)
			self.converted_file.save(self.file + ".pdf")

		except RuntimeError:
			pass

	def __pdf_to_txt__(self) -> None:
		try:
			self.file = filedialog.askopenfilename(title="convert pdf file", filetypes=[("Pdf file (*.pdf)", "*.pdf")], defaultextension=[("Pdf file (*.pdf)", "*.pdf")])
			self.converted_file = Document(self.file)
			self.converted_file.save(self.file + ".txt")

		except RuntimeError:
			pass
		
	def __txt_to_pdf__(self) -> None:
		try:
			self.file = filedialog.askopenfilename(title="convert txt file", filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")])
			self.converted_file = Document(self.file)
			self.converted_file.save(self.file + ".pdf")

		except RuntimeError:
			pass

	def __docx_to_txt__(self) -> None:
		try:
			self.file = filedialog.askopenfilename(title="convert docx file", filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")])
			self.converted_file = Document(self.file)
			self.converted_file.save(self.file + ".txt")
		
		except RuntimeError:
			pass

	def __txt_to_docx__(self) -> None:
		try:
			self.file = filedialog.askopenfilename(title="convert txt file", filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")])
			self.converted_file = Document(self.file)
			self.converted_file.save(self.file + ".docx")

		except RuntimeError:
			pass

	def __code_editor__(self) -> None:
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_open_powershell_button.grid(column=3, row=0)
		self.main_screen_open_terminal_button.grid(column=4, row=0)
		self.main_screen_save_code_as_button.grid(column=5, row=0)
		self.main_screen_open_code_button.grid(column=6, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()
		self.main_screen_word_counter.place_forget()

	def __open_powershell__(self) -> None:
		subprocess.run(["powershell.exe"])

	def __open_shell__(self) -> None:
		subprocess.run(["cmd.exe"])

	def __save_code_as__(self) -> None:
		try:
			with open(filedialog.asksaveasfilename(filetypes=[("Python file (*.py)", "*.py")], defaultextension=[("Python file (*.py)", "*.py")]), "w+", encoding="UTF-8") as self.script_file:
				self.script_file_data = self.main_screen_frame_textbox.get("1.0", END)
				self.script_file.write(self.script_file_data)
				self.script_file.close()
		
		except FileNotFoundError:
			pass

	def __open_code__(self) -> None:
		try:
			with open(filedialog.askopenfilename(title="open script file", filetypes=[("Python file (*.py)", "*.py")], defaultextension=[("Python file (*.py)", "*.py")]), "r+", encoding="UTF-8") as self.openned_script_file:
				self.main_screen_frame_textbox.insert("1.0", self.openned_script_file.read())

		except FileNotFoundError:
			pass
		
	def __settings__(self) -> None:
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_settings_text.place(x=3, y=21)
		self.main_screen_settings_language_text.place(x=3, y=102)
		self.main_screen_settings_language_option.place(x=18, y=142)
		self.main_screen_settings_autosave_text.place(x=3, y=182)
		self.main_screen_settings_autosave_switch.place(x=18, y=222)
		self.main_screen_settings_theme_mode_text.place(x=3, y=262)
		self.main_screen_settings_theme_mode_option.place(x=18, y=302)
		self.main_screen_settings_customatization_text.place(x=18, y=342)
		self.main_screen_settings_customatization_table.place(x=3, y=382)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

		self.main_screen_undo_button.grid_forget()
		self.main_screen_redo_button.grid_forget()
		self.main_screen_word_counter.place_forget()
		self.main_screen_frame.place_forget()

	def __language_settings__(self, pickle) -> None:
		self.main_screen_settings_language_option_variable = self.main_screen_settings_language_option.get()
		with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)

		if self.main_screen_settings_language_option_variable == "Српски":
			showwarning(title="Пажња", message="Рестартуј програм")

		elif self.main_screen_settings_language_option_variable == "English":
			showwarning(title="Warning", message="Restart program")

		else:
			showwarning(title="Внимание", message="Перезагрузите программу")

	def __autosave_settings__(self) -> None:
		self.main_screen_settings_autosave_switch_value = self.main_screen_settings_autosave_value.get()
		with open("my_diary_autosave_settings.pickle", "wb") as self.autosave_data:
			dump(self.main_screen_settings_autosave_switch_value, self.autosave_data)

			if self.language_data == "Српски":
				showwarning(title="Пажња", message="Рестартуј програм")

			elif self.language_data == "English":
				showwarning(title="Warning", message="Restart program")

			else:
				showwarning(title="Внимание", message="Перезагрузите программу")

	def __text_autosave__(self, event) -> None:
		self.main_screen_frame_textbox_text_data = self.main_screen_frame_textbox.get("1.0", END)
		with open('my_diary_saved_text.pickle', 'wb') as self.text_data:
			dump(self.main_screen_frame_textbox_text_data, self.text_data)

	def __theme_settings__(self, pickle) -> None:
		self.main_screen_settings_theme_mode_option_data = self.main_screen_settings_theme_mode_option.get()
		with open("my_diary_theme_settings.pickle", "wb") as self.theme_data:
			dump(self.main_screen_settings_theme_mode_option_data, self.theme_data)

			if self.language_data == "Српски":
				showwarning(title="Пажња", message="Рестартуј програм")

			elif self.language_data == "English":
				showwarning(title="Warning", message="Restart program")

			else:
				showwarning(title="Внимание", message="Перезагрузите программу")

	def __change_text_color__(self, pickle) -> None:
		self.main_screen_settings_customatization_text_color_option_data = self.main_screen_settings_customatization_text_color_option.get()
		with open("my_diary_text_color.pickle", "wb") as self.text_color_data:
			dump(self.main_screen_settings_customatization_text_color_option_data, self.text_color_data)

			if self.language_data == "Српски":
				showwarning(title="Пажња", message="Рестартуј програм")

			elif self.language_data == "English":
				showwarning(title="Warning", message="Restart program")

			else:
				showwarning(title="Внимание", message="Перезагрузите программу")

	def __change_text_field_color__(self, pickle) -> None:
		self.main_screen_settings_customatization_text_field_color_option_data = self.main_screen_settings_customatization_text_field_color_option.get()
		with open("my_diary_text_field_color.pickle", "wb") as self.text_field_color_data:
			dump(self.main_screen_settings_customatization_text_field_color_option_data, self.text_field_color_data)

			if self.language_data == "Српски":
				showwarning(title="Пажња", message="Рестартуј програм")

			elif self.language_data == "English":
				showwarning(title="Warning", message="Restart program")

			else:
				showwarning(title="Внимание", message="Перезагрузите программу")

	def __change_text_field_text_height__(self, pickle) -> None:
		self.main_screen_settings_customatization_text_field_height_option_data = self.main_screen_settings_customatization_text_field_height_option.get()
		with open("my_diary_text_field_text_height.pickle", "wb") as self.text_field_text_height_data:
			dump(self.main_screen_settings_customatization_text_field_height_option_data, self.text_field_text_height_data)

			if self.language_data == "Српски":
				showwarning(title="Пажња", message="Рестартуј програм")

			elif self.language_data == "English":
				showwarning(title="Warning", message="Restart program")

			else:
				showwarning(title="Внимание", message="Перезагрузите программу")

	def __change_button_color__(self, picle) -> None:
		self.main_screen_settings_customatization_button_color_option_data = self.main_screen_settings_customatization_button_color_option.get()
		with open("my_diary_button_color.pickle", "wb") as self.button_color_data:
			dump(self.main_screen_settings_customatization_button_color_option_data, self.button_color_data)

			if self.language_data == "Српски":
				showwarning(title="Пажња", message="Рестартуј програм")

			elif self.language_data == "English":
				showwarning(title="Warning", message="Restart program")

			else:
				showwarning(title="Внимание", message="Перезагрузите программу")

	def __word_count__(self, event) -> None:
		self.main_screen_frame_textbox_data = self.main_screen_frame_textbox.get("0.0", END)
		self.main_screen_word_counter_data_variable.set(value=len(self.main_screen_frame_textbox_data) - 1)
		
	def __program_version__(self, event):
		if self.language_data == "Српски":
			showinfo(title="Version_info", message="версија: 5.0")

		elif self.language_data == "English":
			showinfo(title="Version_info", message="version: 5.0")

		else:
			showinfo(title="Version_info", message="версия: 5.0")

	def __fullscreen__(self, event) -> None:
		self.main_screen_fullscreen_numbers += 1
		if self.main_screen_fullscreen_numbers % 2 == 0:
			self.attributes('-fullscreen', True)
			
			self.main_screen_frame.configure(height=840.55)
			self.main_screen_frame_textbox.configure(height=839)
		
		else:
			self.attributes('-fullscreen', False)

			self.main_screen_frame.configure(height=769)
			self.main_screen_frame_textbox.configure(height=767.5)

	def __open_terminal__(self, event) -> None:
		self.terminal_frame = Terminal(master=self)

	def __html_script__(self, event) -> None:
		self.main_screen_frame_textbox.insert("0.0", f"<!DOCTYPE html> \n  \n  <html lang='en'> \n <head> \n <meta charset='utf-8' /> \n <title></title> \n </head> \n <body> \n \n </body> \n </html>")

	def __right_click_menu__(self, event)-> None:
		try:
			self.main_screen_right_click_menu.tk_popup(event.x_root, event.y_root)

		finally:
			self.main_screen_right_click_menu.grab_release()

	def __copy__(self) -> str:
		self.main_screen_frame_textbox_text_data = self.main_screen_frame_textbox.selection_get()
		return self.main_screen_frame_textbox_text_data

	def __paste__(self) -> None:
		self.main_screen_frame_textbox.insert(self.main_screen_frame_textbox.index("insert"), str(self.main_screen_frame_textbox_text_data))

	def __cut__(self) -> None:
		self.main_screen_frame_textbox.delete(self.main_screen_frame_textbox.index("sel.first"), self.main_screen_frame_textbox.index("sel.last"))	  																																			   

	def __exit__(self) -> None:
		if self.language_data == "Српски":
			self.main_screen_exit = askyesno(title="излаз", message="желите да изађете?")
			if self.main_screen_exit: sys.exit()

		elif self.language_data == "English":
			self.main_screen_exit = askyesno(title="exit", message="would you like to exit?")
			if self.main_screen_exit: sys.exit()

		else:
			self.main_screen_exit = askyesno(title="выход", message="желайте выйти?")
			if self.main_screen_exit: sys.exit()

class Program_taskbar(CTkFrame):
	def __init__(self, master=any, width=300, height=791, border_width=0, corner_radius=5, fg_color="transparent") -> None:
		CTkFrame.__init__(self, master, width, height, corner_radius, border_width, fg_color)

class Terminal(CTkToplevel):
	def __init__(self, master, *args, **kwargs) -> None:
		CTkToplevel.__init__(self, master, *args, **kwargs)

		self.geometry("374x375")
		self.attributes("-toolwindow", True)
		self.resizable(False, False)
		self.title("My Diary command prompt")

		self.terminal_frame = CTkFrame(master=self, height=300, width=300, border_width=1, border_color="white")
		self.terminal_frame.place(x=0, y=0)

		self.terminal_textbox = CTkTextbox(master=self.terminal_frame, height=264, width=296, fg_color="black", text_color="green")
		self.terminal_textbox.place(x=1.5, y=1)

		self.terminal_textbox.insert("0.0", ">>>")
		self.terminal_textbox.configure(state="disabled")

		self.terminal_entry = CTkEntry(master=self.terminal_frame, height=30, width=296, fg_color="black", text_color="green")
		self.terminal_entry.place(x=1.5, y=267)

		self.terminal_entry.bind('<Return>', self.__terminal_action__)

	def __terminal_action__(self, event) -> None:
		self.terminal_entry_data = self.terminal_entry.get()

		self.terminal_textbox.configure(state="normal")
		if self.terminal_entry_data == "clear text":
			program.__clear_text__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open taskbar":
			program.__open_taskbar__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "close taskbar":
			program.__close_taskbar__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "quit operation":
			program.__quit_operation__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)
		
		elif self.terminal_entry_data == "save text as ...":
			program.__save_text_as__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as docx":
			program.__save_text_as_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "clear text":
			program.__clear_text__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open ... file":
			program.__open_file__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open docx file":
			program.__open_file_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert pdf to docx":
			program.__pdf_to_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert docx to pdf":
			program.__docx_to_pdf__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert pdf to txt":
			program.__pdf_to_txt__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert txt to pdf":
			program.__txt_to_pdf__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert docx to txt":
			program.__docx_to_txt__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert txt to docx":
			program.__txt_to_docx__()
			
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "run code":
			program.__run_code__()
			
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open shell":
			program.__open_shell__()
			
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open powershell":
			program.__open_powershell__()
			
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save code as":
			program.__save_code_as__()
			
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open code":
			program.__open_code__()
			
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "programm version":
			program.__programm_version__(event)

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "html script":
			program.__html_script__(event)

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "help":
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, f"commands: \n 1) save text / code as ... \n 2) open ... file / code / shell / powershell \n 3) clear text \n 4) convert ... to ... \n 5) exit program / terminal \n 6) open/close taskbar \n 7) quit operation \n 8) html script \n 9) programm version \n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "exit terminal":
			self.destroy()

		elif self.terminal_entry_data == "exit program":
			program.__exit__()
		
		else:
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + " - command doesn't exist" + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

if __name__ == "__main__":
	try:
		program = Program()
		program.mainloop()
	
	except FileNotFoundError:
		showerror(message="error: missing data file")

	except TclError:
		showerror(message="error: missing icon file")

else:
	showwarning(title="not script file", message="this is not script file")