from customtkinter import * 
from tkinter import * 
from pickle import *
from tkinter.messagebox import *
from aspose.words import *
from tkinter import filedialog
from docx.shared import RGBColor, Pt
import os, sys, docx

class Program(CTk):
	def __init__(self, *args, **kwargs) -> CTk:
		CTk.__init__(self, *args, **kwargs)

		set_widget_scaling(1.251)
		set_default_color_theme("dark-blue")
		deactivate_automatic_dpi_awareness()

		self.title("My diary")
		self.iconbitmap("my diary icon.ico")

		self.bind('<F1>', self.__program_version__)  
		self.bind('<Control_L>' + '<f>', self.__fullscreen__) 
		self.bind('<Alt_L>' + '<F4>', self.__exit__)
		self.bind('<Alt_L>' + '<t>', self.__open_terminal__)
		self.bind('<Escape>', self.__close_taskbar_with_peripherals__)
		self.bind('<Control_L>' + '<t>', self.__open_taskbar_with_keybord__)
		self.bind('<Button-3>', self.__right_click_menu__)

		self.protocol("WM_DELETE_WINDOW", self.__exit__)
		
		self.main_screen_fullscreen_numbers = 1

		with open("my_diary_settings.pickle", "rb") as self.data:
			self.new_data = load(self.data)

		with open('my_diary_saved_text.pickle', "rb") as self.text_data:
			self.new_text_data = load(self.text_data)

		with open("my_diary_autosave_settings.pickle", "rb") as self.autosave_data:
			self.new_autosaved_data = load(self.autosave_data)

		with open("my_diary_theme_settings.pickle", "rb") as self.theme_data:
			self.new_theme_data = load(self.theme_data)

		self.main_screen_taskbar_button = CTkButton(master=self, text="☰", height=20, width=10, corner_radius=0, fg_color="green", command=self.__open_taskbar__)
		self.main_screen_taskbar_button.grid(column=0, row=0)

		self.main_screen_frame = CTkFrame(master=self, height=769, width=1535, border_width=1, border_color=("black", "white"), corner_radius=0)
		self.main_screen_frame.place(x=0, y=22.35)

		self.main_screen_frame_textbox = CTkTextbox(master=self.main_screen_frame, height=767.5, width=1533.57, corner_radius=0, undo=True)
		self.main_screen_frame_textbox.place(x=1, y=1)

		self.main_screen_frame_textbox.bind('<KeyRelease>', self.__word_count__)
		self.main_screen_frame_textbox.bind('<F2>', self.__html_script__)
		self.main_screen_frame_textbox.bind('<KeyPress>', self.__close_taskbar_with_peripherals__)
		self.main_screen_frame_textbox.bind('<ButtonRelease>', self.__close_taskbar_with_peripherals__)

		self.main_screen_taskbar = Program_taskbar(master=self)

		self.main_screen_taskbar_exit_button = CTkButton(master=self.main_screen_taskbar, text="↵", height=20, width=10, corner_radius=5, fg_color="green", command=self.__close_taskbar__)
		self.main_screen_taskbar_exit_button.place(x=278, y=2)

		self.main_screen_undo_button = CTkButton(master=self.main_screen_taskbar, text="⟲", height=20, width=10, corner_radius=5, fg_color="green", command=lambda: self.main_screen_frame_textbox.edit_undo())
		self.main_screen_undo_button.place(x=2, y=2)

		self.main_screen_redo_button = CTkButton(master=self.main_screen_taskbar, text="⟳", height=20, width=10, corner_radius=5, fg_color="green", command=lambda: self.main_screen_frame_textbox.edit_redo())
		self.main_screen_redo_button.place(x=32, y=2)

		self.main_screen_save_button = CTkButton(master=self.main_screen_taskbar, text="сачувај текст", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__save__)
		self.main_screen_save_button.place(x=2, y=32)

		self.main_screen_save_as_txt_button = CTkButton(master=self, text="сачувај као txt", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_txt__)

		self.main_screen_save_as_docx_button = CTkButton(master=self, text="сачувај као docx", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_docx__)

		self.main_screen_save_as_py_button = CTkButton(master=self, text="сачувај као python", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_py__)

		self.main_screen_save_as_java_button = CTkButton(master=self, text="сачувај као java", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_java__)

		self.main_screen_save_as_kotlin_button = CTkButton(master=self, text="сачувај као kotlin", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_kotlin__)

		self.main_screen_save_as_html_button = CTkButton(master=self, text="сачувај као html", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_html__)

		self.main_screen_clear_button = CTkButton(master=self.main_screen_taskbar, text="обриши текст", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=lambda: self.main_screen_frame_textbox.delete("1.0", END))
		self.main_screen_clear_button.place(x=2, y=67)

		self.main_screen_edit_text_button = CTkButton(master=self.main_screen_taskbar, text="уреди текст", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__edit_text_)
		self.main_screen_edit_text_button.place(x=2, y=102)

		self.main_screen_edit_font_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["System", "Terminal", "Ubuntu", "Script", "Roman", "Modern", "MS Serif"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_size_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["1", "2", "4", "5", "6", "8", "11", "12", "13", "14", "16", "18", "20", "22", "24", "26", "28", "30", "32", "34", "36", "38", "40", "42", "44", "46", "48", "50", "60", "70", "80", "90", "100"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_color_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["black", "white", "red", "green", "blue"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_slant_button =  CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["roman", "italic"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_weight_button =  CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["normal", "bold"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_edit_underline_button =  CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["not underlined", "underlined"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white", command=self.__edit_text_font__)

		self.main_screen_open_file_button = CTkButton(master=self.main_screen_taskbar, text="отвори фајл", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__open_file__)
		self.main_screen_open_file_button.place(x=2, y=137)
	
		self.main_screen_open_file_txt_button = CTkButton(master=self, text="отвори txt фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_txt__)
	
		self.main_screen_open_file_docx_button = CTkButton(master=self, text="отвори docx фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_docx__)
	
		self.main_screen_open_file_py_button = CTkButton(master=self, text="отвори python фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_py__)

		self.main_screen_open_file_html_button = CTkButton(master=self, text="отвори html фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_html__)

		self.main_screen_coding_button = CTkButton(master=self.main_screen_taskbar, text="кодирање", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__coding__)
		self.main_screen_coding_button.place(x=2, y=172)

		self.main_screen_coding_binar_values = {"a": "0", "b": "1", "c": "10", "d": "11", "e": "100", "f": "101", "g": "110", "h": "111", "i": "1000", "j": "1001", "k": "1010", "l": "1011", "m": "1100", "n": "1101", "o": "1110", "p": "1111", "q": "10000", "r": "10001", "s": "10010", "t": "10011", "u": "10100", "v": "10101", "w": "10111", "x": "11000", "y": "11001", "z": "11011", " ": " ", ".": "."}

		self.main_screen_coding_ceaser_values = {"a": "d", "b": "e", "c": "f", "d": "g", "e": "h", "f": "i", "g": "j", "h": "k", "i": "l", "j": "m", "k": "n", "l": "o", "m": "p", "n": "q", "o": "r", "p": "s", "q": "t", "r": "u", "s": "v", "t": "w", "u": "x", "v": "y", "w": "z", "x": "a", "y": "b", "z": "c", " ": " ", ".": "."}
		
		self.main_screen_coding_fibonachi_values = {"a": "0", "b": "1", "c": "2", "d": "3", "e": "5", "f": "8", "g": "13", "h": "21", "i": "34", "j": "55", "k": "89", "l": "144", "m": "233", "n": "377", "o": "610", "p": "987", "q": "1597", "r": "2584", "s": "4181", "t": "6765", "u": "10946", "v": "17711", "w": "28657", "x": "46368", "y": "75025", "z": "121393", " ": " ", ".": "."}

		self.main_screen_coding_numerical_values = {"a": "1", "b": "2", "c" : "3", "d": "4", "e": "5", "f": "6", "g": "7", "h": "8", "i": "9", "j": "10", "k": "11", "l": "12", "m": "13", "q": "14", "r": "15", "s": "16", "t": "17", "u": "18", "v": "19", "w": "20", "x": "21", "y": "22", "z": "23", " ": " ", ".": "."}

		self.main_screen_coding_mode_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["binar", "ceaser", "fibonachi", "numerical"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_text_color="white")

		self.main_screen_coding_function_button = CTkButton(master=self, text="кодирај", height=20, width=25, corner_radius=0, fg_color="green", command=self.__message_encoding__)

		self.main_screen_converter_button = CTkButton(master=self.main_screen_taskbar, text="претварач", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__convert__)
		self.main_screen_converter_button.place(x=2, y=207)

		self.main_screen_pdf_to_word_converter_button = CTkButton(master=self, text="из Pdf у docx", height=20, width=25, corner_radius=0, fg_color="green", command=self.__pdf_to_docx__)

		self.main_screen_word_to_pdf_converter_button = CTkButton(master=self, text="из docx у Pdf", height=20, width=25, corner_radius=0, fg_color="green", command=self.__docx_to_pdf__)

		self.main_screen_pdf_to_txt_converter_button = CTkButton(master=self, text="из Pdf у txt", height=20, width=25, corner_radius=0, fg_color="green", command=self.__pdf_to_txt__)

		self.main_screen_txt_to_pdf_converter_button = CTkButton(master=self, text="из txt у Pdf", height=20, width=25, corner_radius=0, fg_color="green", command=self.__txt_to_pdf__)

		self.main_screen_word_to_txt_converter_button = CTkButton(master=self, text="из docx y txt", height=20, width=25, corner_radius=0, fg_color="green", command=self.__docx_to_txt__)

		self.main_screen_txt_to_word_converter_button = CTkButton(master=self, text="из txt у docx", height=20, width=25, corner_radius=0, fg_color="green", command=self.__txt_to_docx__)

		self.main_screen_right_click_menu = Menu(self, tearoff=0)

		self.main_screen_right_click_save_menu = Menu(self, tearoff=0)

		self.main_screen_right_click_open_menu = Menu(self, tearoff=0)

		self.main_screen_right_click_convert_menu = Menu(self, tearoff=0)

		self.main_screen_settings_button = CTkButton(master=self.main_screen_taskbar, text="подешавања", height=20, width=295, corner_radius=5, fg_color="green", font=("Roman", 22), command=self.__settings__)
		self.main_screen_settings_button.place(x=2, y=242)

		self.main_screen_settings_text = CTkLabel(master=self, text="подешавања", font=("Roboto Bold", 72))

		self.main_screen_settings_language_text = CTkLabel(master=self, text="језици", font=("Roboto Bold", 36))

		self.main_screen_settings_language_option = CTkSegmentedButton(master=self, values=["Српски", "English", "Русский"], selected_color="green", command=self.__language_settings__)

		self.language_data = str(self.new_data)

		self.main_screen_settings_language_option.set(self.language_data)

		self.main_screen_settings_autosave_text = CTkLabel(master=self, text="ауточување", font=("Roboto Bold", 36))

		self.main_screen_settings_autosave_value = StringVar(value=self.new_autosaved_data)

		self.main_screen_settings_autosave_switch_variable = str(self.new_autosaved_data)

		self.main_screen_settings_autosave_switch = CTkSwitch(master=self, text="ауточување", variable=self.main_screen_settings_autosave_value,onvalue="on", offvalue="off", progress_color="green", command=self.__autosave_settings__)

		self.autosaved_text = str(self.new_text_data)

		self.main_screen_settings_autosave_switch_value = self.main_screen_settings_autosave_value.get()
			
		self.main_screen_settings_theme_mode_text = CTkLabel(master=self, text="тема", font=("Roboto Bold", 36))

		self.main_screen_settings_theme_mode_option = CTkSegmentedButton(master=self, values=["System", "Dark", "Light"], selected_color="green", command=self.__theme_settings__)

		self.theme = str(self.new_theme_data)

		self.main_screen_settings_theme_mode_option.set(self.theme)
		
		if self.language_data == "Српски":
			self.main_screen_save_button.configure(text="сачувај текст")
			self.main_screen_save_as_txt_button.configure(text="сачувај као txt")
			self.main_screen_save_as_docx_button.configure(text="сачувај као docx")
			self.main_screen_save_as_py_button.configure(text="сачувај као python")
			self.main_screen_save_as_java_button.configure(text="сачувај као java")
			self.main_screen_save_as_kotlin_button.configure(text="сачувај као kotlin")
			self.main_screen_save_as_html_button.configure(text="сачувај као html")

			self.main_screen_clear_button.configure(text="обриши текст")
			self.main_screen_edit_text_button.configure(text="уреди текст")
			
			self.main_screen_open_file_button.configure(text="отвори фајл")
			self.main_screen_open_file_txt_button.configure(text="отвори txt фајл")
			self.main_screen_open_file_docx_button.configure(text="отвори docx фајл")
			self.main_screen_open_file_py_button.configure(text="отвори python фајл")
			self.main_screen_open_file_html_button.configure(text="отвори html фајл")

			self.main_screen_coding_button.configure(text="кодирање")
			self.main_screen_coding_function_button.configure(text="кодирај")

			self.main_screen_converter_button.configure(text="претварач")
			self.main_screen_pdf_to_word_converter_button.configure(text="из Pdf у docx")
			self.main_screen_word_to_pdf_converter_button.configure(text="из docx у Pdf")
			self.main_screen_pdf_to_txt_converter_button.configure(text="из Pdf у txt")
			self.main_screen_txt_to_pdf_converter_button.configure(text="из txt у Pdf")
			self.main_screen_word_to_txt_converter_button.configure(text="из docx у txt")
			self.main_screen_txt_to_word_converter_button.configure(text="из txt у docx")

			self.main_screen_settings_button.configure(text="подешавања")
			self.main_screen_settings_text.configure(text="подешавања")
			self.main_screen_settings_language_text.configure(text="језици")
			self.main_screen_settings_autosave_text.configure(text="ауточување")
			self.main_screen_settings_autosave_switch.configure(text="ауточување")
			self.main_screen_settings_theme_mode_text.configure(text="теме")

			self.main_screen_right_click_menu.add_cascade(label="сачувај као", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_save_menu)
			self.main_screen_right_click_save_menu.add_command(label=".txt", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_txt__)
			self.main_screen_right_click_save_menu.add_command(label=".docx", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_docx__)
			self.main_screen_right_click_save_menu.add_command(label=".py", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_py__)
			self.main_screen_right_click_save_menu.add_command(label=".java", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_java__)
			self.main_screen_right_click_save_menu.add_command(label=".kt", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_kotlin__)
			self.main_screen_right_click_save_menu.add_command(label=".html", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_html__)

			self.main_screen_right_click_menu.add_cascade(label="отвори као", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_open_menu)
			self.main_screen_right_click_open_menu.add_command(label="текст", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_txt__)
			self.main_screen_right_click_open_menu.add_command(label="ворд", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_docx__)
			self.main_screen_right_click_open_menu.add_command(label="python", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_py__)
			self.main_screen_right_click_open_menu.add_command(label="html", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_html__)

			self.main_screen_right_click_menu.add_cascade(label="конвертирај", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_convert_menu)
			self.main_screen_right_click_convert_menu.add_command(label="из docx у pdf", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="из pdf у docx", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="из txt у pdf", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="из pdf у txt", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_txt__)
			self.main_screen_right_click_convert_menu.add_command(label="из txt у docx", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="из docx у txt", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_txt__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="обриши текст", font=("Roman", 16), background="green", foreground="white", command=lambda: self.main_screen_frame_textbox.delete("1.0", END))

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="копирај", font=("Roman", 16), background="green", foreground="white", command=self.__copy__)
			self.main_screen_right_click_menu.add_command(label="залепи", font=("Roman", 16), background="green", foreground="white", command=self.__paste__)
			self.main_screen_right_click_menu.add_command(label="исеци", font=("Roman", 16), background="green", foreground="white", command=self.__cut__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="изаћи", font=("Roman", 16), background="green", foreground="white", command=self.__exit__)

		elif self.language_data == "English":
			self.main_screen_save_button.configure(text="save text")
			self.main_screen_save_as_txt_button.configure(text="save as txt")
			self.main_screen_save_as_docx_button.configure(text="save as docx")
			self.main_screen_save_as_py_button.configure(text="save as python")
			self.main_screen_save_as_java_button.configure(text="save as java")
			self.main_screen_save_as_kotlin_button.configure(text="save as kotlin")
			self.main_screen_save_as_html_button.configure(text="save as html")

			self.main_screen_clear_button.configure(text="clear text")
			self.main_screen_edit_text_button.configure(text="edit text")
			
			self.main_screen_open_file_button.configure(text="open file")
			self.main_screen_open_file_txt_button.configure(text="open txt file")
			self.main_screen_open_file_docx_button.configure(text="open docx file")
			self.main_screen_open_file_py_button.configure(text="open python file")
			self.main_screen_open_file_html_button.configure(text="open html file")

			self.main_screen_coding_button.configure(text="coding")
			self.main_screen_coding_function_button.configure(text="code")

			self.main_screen_converter_button.configure(text="converterer")
			self.main_screen_pdf_to_word_converter_button.configure(text="from Pdf to docx")
			self.main_screen_word_to_pdf_converter_button.configure(text="from docx to Pdf")
			self.main_screen_pdf_to_txt_converter_button.configure(text="from Pdf to txt")
			self.main_screen_txt_to_pdf_converter_button.configure(text="from txt to Pdf")
			self.main_screen_word_to_txt_converter_button.configure(text="from docx to txt")
			self.main_screen_txt_to_word_converter_button.configure(text="from txt to docx")

			self.main_screen_settings_button.configure(text="settings")
			self.main_screen_settings_text.configure(text="settings")
			self.main_screen_settings_language_text.configure(text="languages")
			self.main_screen_settings_autosave_text.configure(text="autosave")
			self.main_screen_settings_autosave_switch.configure(text="autosave")
			self.main_screen_settings_theme_mode_text.configure(text="theme")

			self.main_screen_right_click_menu.add_cascade(label="save as", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_save_menu)
			self.main_screen_right_click_save_menu.add_command(label=".txt", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_txt__)
			self.main_screen_right_click_save_menu.add_command(label=".docx", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_docx__)
			self.main_screen_right_click_save_menu.add_command(label=".py", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_py__)
			self.main_screen_right_click_save_menu.add_command(label=".java", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_java__)
			self.main_screen_right_click_save_menu.add_command(label=".kt", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_kotlin__)
			self.main_screen_right_click_save_menu.add_command(label=".html", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_html__)

			self.main_screen_right_click_menu.add_cascade(label="open as", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_open_menu)
			self.main_screen_right_click_open_menu.add_command(label="text", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_txt__)
			self.main_screen_right_click_open_menu.add_command(label="word", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_docx__)
			self.main_screen_right_click_open_menu.add_command(label="python", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_py__)
			self.main_screen_right_click_open_menu.add_command(label="html", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_html__)

			self.main_screen_right_click_menu.add_cascade(label="convert", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_convert_menu)
			self.main_screen_right_click_convert_menu.add_command(label="from docx to pdf", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="from pdf to docx", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="from txt to pdf", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="from pdf to txt", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_txt__)
			self.main_screen_right_click_convert_menu.add_command(label="from txt to docx", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="from docx to txt", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_txt__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="delete text", font=("Roman", 16), background="green", foreground="white", command=lambda: self.main_screen_frame_textbox.delete("1.0", END))

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="copy", font=("Roman", 16), background="green", foreground="white", command=self.__copy__)
			self.main_screen_right_click_menu.add_command(label="paste", font=("Roman", 16), background="green", foreground="white", command=self.__paste__)
			self.main_screen_right_click_menu.add_command(label="cut", font=("Roman", 16), background="green", foreground="white", command=self.__cut__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="exit", font=("Roman", 16), background="green", foreground="white", command=self.__exit__)

		elif self.language_data == "Русский":
			self.main_screen_save_button.configure(text="сохранить текст")
			self.main_screen_save_as_txt_button.configure(text="сохранить как текст")
			self.main_screen_save_as_docx_button.configure(text="сохранить как ворд")
			self.main_screen_save_as_py_button.configure(text="сохранить как python")
			self.main_screen_save_as_java_button.configure(text="сохранить как java")
			self.main_screen_save_as_kotlin_button.configure(text="сохранить как kotlin")
			self.main_screen_save_as_html_button.configure(text="сохранить как html")

			self.main_screen_clear_button.configure(text="удалить текст")
			self.main_screen_edit_text_button.configure(text="редактировать текст")
			
			self.main_screen_open_file_button.configure(text="открыть файл")
			self.main_screen_open_file_txt_button.configure(text="открыть текстовый файл")
			self.main_screen_open_file_docx_button.configure(text="открыть ворд файл")
			self.main_screen_open_file_py_button.configure(text="открыть python файл")
			self.main_screen_open_file_html_button.configure(text="открыть html файл")
			
			self.main_screen_coding_button.configure(text="кодировка")
			self.main_screen_coding_function_button.configure(text="кодировать")

			self.main_screen_converter_button.configure(text="конвертер")
			self.main_screen_pdf_to_word_converter_button.configure(text="из Пдф в ворд")
			self.main_screen_word_to_pdf_converter_button.configure(text="из ворд в Пдф")
			self.main_screen_pdf_to_txt_converter_button.configure(text="из Пдф в текстовый файл")
			self.main_screen_txt_to_pdf_converter_button.configure(text="из текстового файла в Пдф")
			self.main_screen_word_to_txt_converter_button.configure(text="из ворд в текстовый файл")
			self.main_screen_txt_to_word_converter_button.configure(text="из текстового файла в ворд")

			self.main_screen_settings_button.configure(text="настройки")
			self.main_screen_settings_text.configure(text="настройки")
			self.main_screen_settings_language_text.configure(text="языки")
			self.main_screen_settings_autosave_text.configure(text="автосохранение")
			self.main_screen_settings_autosave_switch.configure(text="автосохранение")
			self.main_screen_settings_theme_mode_text.configure(text="темы")

			self.main_screen_right_click_menu.add_cascade(label="сохранить как", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_save_menu)
			self.main_screen_right_click_save_menu.add_command(label="текст", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_txt__)
			self.main_screen_right_click_save_menu.add_command(label="ворд", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_docx__)
			self.main_screen_right_click_save_menu.add_command(label=".py", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_py__)
			self.main_screen_right_click_save_menu.add_command(label=".java", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_java__)
			self.main_screen_right_click_save_menu.add_command(label=".kt", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_kotlin__)
			self.main_screen_right_click_save_menu.add_command(label=".html", font=("Roman", 16), background="green", foreground="white", command=self.__save_text_as_html__)

			self.main_screen_right_click_menu.add_cascade(label="открыть как", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_open_menu)
			self.main_screen_right_click_open_menu.add_command(label="текст", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_txt__)
			self.main_screen_right_click_open_menu.add_command(label="ворд", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_docx__)
			self.main_screen_right_click_open_menu.add_command(label="python", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_py__)
			self.main_screen_right_click_open_menu.add_command(label="html", font=("Roman", 16), background="green", foreground="white", command=self.__open_file_html__)

			self.main_screen_right_click_menu.add_cascade(label="конвертировать", font=("Roman", 16), background="green", foreground="white", menu=self.main_screen_right_click_convert_menu)
			self.main_screen_right_click_convert_menu.add_command(label="из ворд в Пдф", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="из Пдф в ворд", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="из текста в Пдф", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_pdf__)
			self.main_screen_right_click_convert_menu.add_command(label="из Пдф в текст", font=("Roman", 16), background="green", foreground="white", command=self.__pdf_to_txt__)
			self.main_screen_right_click_convert_menu.add_command(label="из текста в ворд", font=("Roman", 16), background="green", foreground="white", command=self.__txt_to_docx__)
			self.main_screen_right_click_convert_menu.add_command(label="из ворда в текст", font=("Roman", 16), background="green", foreground="white", command=self.__docx_to_txt__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="удалить текст", font=("Roman", 16), background="green", foreground="white", command=lambda: self.main_screen_frame_textbox.delete("1.0", END))

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="копировать", font=("Roman", 16), background="green", foreground="white", command=self.__copy__)
			self.main_screen_right_click_menu.add_command(label="вставить", font=("Roman", 16), background="green", foreground="white", command=self.__paste__)
			self.main_screen_right_click_menu.add_command(label="вырезать", font=("Roman", 16), background="green", foreground="white", command=self.__cut__)

			self.main_screen_right_click_menu.add_separator()

			self.main_screen_right_click_menu.add_command(label="выход", font=("Roman", 16), background="green", foreground="white", command=self.__exit__)

		if self.theme == "System":
			set_appearance_mode("system")

		elif self.theme == "Dark":
			set_appearance_mode("dark")

		elif self.theme == "Light":
			set_appearance_mode("light")


		if self.main_screen_settings_autosave_switch_variable == "on":
			self.main_screen_frame_textbox.insert("1.0", self.autosaved_text)

			self.bind('<Return>', self.__text_autosave__)
			self.bind('<Control_L>' + '<s>', self.__text_autosave__)

		elif self.main_screen_settings_autosave_switch_variable == "off":
			self.unbind('<Return>')
			self.unbind('<Control_L>' + '<s>')

		self.main_screen_word_counter_variable = 0

		self.main_screen_word_counter_data_variable = IntVar(value=self.main_screen_word_counter_variable)

		self.main_screen_word_counter = CTkLabel(master=self, width=25, height=20, corner_radius=0, textvariable=self.main_screen_word_counter_data_variable)
		self.main_screen_word_counter.place(x=1500, y=0)
				  
		self.main_screen_coding_frame = CTkFrame(master=self, height=769, width=807, border_width=1, border_color=("black", "white"), corner_radius=0)

		self.main_screen_coding_frame_textbox = CTkTextbox(master=self.main_screen_coding_frame, height=767.5, width=804, corner_radius=0)

	def __open_taskbar__(self):
		self.main_screen_taskbar.grid(column=0, row=0)
		self.main_screen_taskbar_button.grid_forget()

	def __open_taskbar_with_keybord__(self, event):
		self.main_screen_taskbar.grid(column=0, row=0)
		self.main_screen_taskbar_button.grid_forget()

	def __close_taskbar__(self):
		self.main_screen_taskbar.grid_forget()
		self.main_screen_taskbar_button.grid(column=0, row=0)

	def __close_taskbar_with_peripherals__(self, event):
		self.main_screen_taskbar.grid_forget()
		self.main_screen_taskbar_button.grid(column=0, row=0)

	def __quit_operation__(self):
		self.bind('<Control_L>' + '<Tab>' + '<t>', self.__open_taskbar_with_keybord__)

		self.main_screen_taskbar_button.configure(text="☰", command=self.__open_taskbar__)

		self.main_screen_save_as_txt_button.grid_forget()
		self.main_screen_save_as_docx_button.grid_forget()
		self.main_screen_save_as_py_button.grid_forget()
		self.main_screen_save_as_java_button.grid_forget()
		self.main_screen_save_as_kotlin_button.grid_forget()
		self.main_screen_save_as_html_button.grid_forget()
		
		self.main_screen_edit_font_button.grid_forget()
		self.main_screen_edit_size_button.grid_forget()
		self.main_screen_edit_color_button.grid_forget()
		self.main_screen_edit_slant_button.grid_forget()
		self.main_screen_edit_weight_button.grid_forget()
		self.main_screen_edit_underline_button.grid_forget()

		self.main_screen_open_file_txt_button.grid_forget()
		self.main_screen_open_file_docx_button.grid_forget()
		self.main_screen_open_file_py_button.grid_forget()
		self.main_screen_open_file_html_button.grid_forget()

		self.main_screen_frame.configure(width=1535)
		self.main_screen_frame_textbox.configure(width=1530)

		self.main_screen_coding_frame.place_forget()
		self.main_screen_coding_frame_textbox.place_forget()
		self.main_screen_coding_mode_button.place_forget()
		self.main_screen_coding_function_button.place_forget()

		self.main_screen_coding_frame_textbox.delete("1.0", END)

		self.main_screen_pdf_to_word_converter_button.grid_forget()
		self.main_screen_word_to_pdf_converter_button.grid_forget()
		self.main_screen_pdf_to_txt_converter_button.grid_forget()
		self.main_screen_txt_to_pdf_converter_button.grid_forget()
		self.main_screen_word_to_txt_converter_button.grid_forget()
		self.main_screen_txt_to_word_converter_button.grid_forget()

		self.main_screen_settings_text.place_forget()
		self.main_screen_settings_language_text.place_forget()
		self.main_screen_settings_language_option.place_forget()						          
		self.main_screen_settings_autosave_text.place_forget()
		self.main_screen_settings_autosave_switch.place_forget()
		self.main_screen_settings_theme_mode_text.place_forget()
		self.main_screen_settings_theme_mode_option.place_forget()

		self.main_screen_frame.place(x=0, y=22.35)
		self.main_screen_word_counter.place(x=1500, y=0)

	def __clear_text__(self):
		self.main_screen_frame_textbox.delete("1.0", END)

	def __save__(self):
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_save_as_txt_button.grid(column=1, row=0)
		self.main_screen_save_as_docx_button.grid(column=2, row=0)
		self.main_screen_save_as_py_button.grid(column=3, row=0)
		self.main_screen_save_as_java_button.grid(column=4, row=0)
		self.main_screen_save_as_kotlin_button.grid(column=5, row=0)
		self.main_screen_save_as_html_button.grid(column=6, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()
					
	def __save_text_as_txt__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")]), "w+", encoding="UTF-8")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_docx__(self):
		try:
			self.file = docx.Document()
			self.file_data = self.main_screen_frame_textbox.get("1.0", END)
			self.file_run = self.file.add_paragraph().add_run(self.file_data)
			self.font = self.file_run.font
			self.font.name = str(self.main_screen_edit_font_button_data)
			self.font.size = Pt(int(self.main_screen_edit_size_button_data))
			if self.main_screen_edit_color_button_data == "black":
				self.font.color.rgb = RGBColor(0, 0, 0)

			elif self.main_screen_edit_color_button_data == "white":
				self.font.color.rgb = RGBColor(255, 255, 255)

			elif self.main_screen_edit_color_button_data == "red":
				self.font.color.rgb = RGBColor(250, 0, 0)

			elif self.main_screen_edit_color_button_data == "green":
				self.font.color.rgb = RGBColor(0, 255, 0)

			elif self.main_screen_edit_color_button_data == "blue":
				self.font.color.rgb = RGBColor(0, 0, 255)

						
			if self.main_screen_edit_slant_button_data == "italic":
				self.file_run.italic = True

			else:
				self.file_run.italic = False


			if self.main_screen_edit_weight_button_data == "bold":
				self.file_run.bold = True

			else:
				self.file_run.bold = False


			if self.main_screen_edit_underline_button_data == "underlined":
				self.file_run.underline = True

			else:
				self.file_run.underline = False

			self.file.save(filedialog.asksaveasfilename(filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")]))

		except AttributeError:
			if self.language_data == "Српски":
				showerror(title="Грешка", message="Појавила се грешка. Нисте редактирали текст. Нисам мого да сачувам фајл")

			elif self.language_data == "English":
				showerror(title="Error", message="An error has occured. Text wasn't edited. Couldn't save the file")

			elif self.language_data == "Русский":
				showerror(title="Ошибка", message="Появилась ошибка. Текст неотредоктирован. Не смог охранить файл")

	def __save_text_as_py__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Python file (*.py)", "*.py")], defaultextension=[("Python file (*.py)", "*.py")]), "w+")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_java__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Java file (*.java)", "*.java")], defaultextension=[("Java file (*.java)", "*.java")]), "w+")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_kotlin__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Kotlin file (*.kt)", "*.kt")], defaultextension=[("Kotlin file (*.kt)", "*.kt")]), "w+")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_html__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("html file (*.html)", "*.html")], defaultextension=[("Python file (*.html)", "*.html")]), "w+")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __edit_text_(self):
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_edit_font_button.grid(column=1, row=0)
		self.main_screen_edit_size_button.grid(column=2, row=0)
		self.main_screen_edit_color_button.grid(column=3, row=0)
		self.main_screen_edit_slant_button.grid(column=4, row=0)
		self.main_screen_edit_weight_button.grid(column=5, row=0)
		self.main_screen_edit_underline_button.grid(column=6, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

	def __edit_text_font__(self, configure):
		self.main_screen_edit_font_button_data = self.main_screen_edit_font_button.get()
		self.main_screen_edit_size_button_data = self.main_screen_edit_size_button.get()
		self.main_screen_edit_color_button_data = self.main_screen_edit_color_button.get()
		self.main_screen_edit_slant_button_data = self.main_screen_edit_slant_button.get()
		self.main_screen_edit_weight_button_data = self.main_screen_edit_weight_button.get()
		self.main_screen_edit_underline_button_data = self.main_screen_edit_underline_button.get()

		if self.main_screen_edit_underline_button_data == "not underlined":		
			self.main_screen_frame_textbox.configure(font=CTkFont(family=self.main_screen_edit_font_button_data, size=int(self.main_screen_edit_size_button_data), slant=self.main_screen_edit_slant_button_data, weight=self.main_screen_edit_weight_button_data, underline=False), text_color=self.main_screen_edit_color_button_data)

		elif self.main_screen_edit_underline_button_data == "underlined":		
			self.main_screen_frame_textbox.configure(font=CTkFont(family=self.main_screen_edit_font_button_data, size=int(self.main_screen_edit_size_button_data), slant=self.main_screen_edit_slant_button_data, weight=self.main_screen_edit_weight_button_data, underline=True), text_color=self.main_screen_edit_color_button_data)
		
	def __open_file__(self):
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_open_file_txt_button.grid(column=1, row=0)
		self.main_screen_open_file_docx_button.grid(column=2, row=0)
		self.main_screen_open_file_py_button.grid(column=3, row=0)
		self.main_screen_open_file_html_button.grid(column=4, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

	def __open_file_txt__(self):
		self.openned_file = open(filedialog.askopenfilename(title="open txt file", filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")]), "r+", encoding="UTF-8")
		self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

	def __open_file_docx__(self):
		self.openned_file_name = filedialog.askopenfilename(title="open docx file", filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")])
		self.openned_file = docx.Document(self.openned_file_name)
		self.openned_file_data = []
		for self.paragraphs in self.openned_file.paragraphs:
			self.openned_file_data.append(self.paragraphs.text)

		self.main_screen_frame_textbox.insert("1.0", "\n".join(self.openned_file_data))

	def __open_file_py__(self):
		self.openned_file = open(filedialog.askopenfilename(title="open python file", filetypes=[("Python file (*.py)", "*.py")], defaultextension=[("Python file (*.py)", "*.py")]), "r+", encoding="UTF-8")
		self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

	def __open_file_html__(self):
		self.openned_file = open(filedialog.askopenfilename(title="open html file", filetypes=[("html file (*.html)", "*.html")], defaultextension=[("html file (*.html)", "*.html")]), "r+", encoding="UTF-8")
		self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

	def __coding__(self):
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_frame.configure(width=715.5)
		self.main_screen_frame_textbox.configure(width=712.75)
		
		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_coding_frame.place(x=729, y=22)
		self.main_screen_coding_frame_textbox.place(x=1, y=1)
		self.main_screen_coding_mode_button.place(x=729, y=0)
		self.main_screen_coding_function_button.place(x=803.5, y=0)

		self.main_screen_word_counter.place_forget()

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

	def __message_encoding__(self):
		self.main_screen_coding_frame_textbox.delete("1.0", END)

		self.main_screen_coding_mode_button_data = self.main_screen_coding_mode_button.get()
		self.main_screen_frame_textbox_data = self.main_screen_frame_textbox.get("1.0", END)
		if self.main_screen_coding_mode_button_data == "binar":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_binar_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_binar_values.get(key), "-1.0")
		
		elif self.main_screen_coding_mode_button_data == "ceaser":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_ceaser_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_ceaser_values.get(key), "-1.0")

		elif self.main_screen_coding_mode_button_data == "fibonachi":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_fibonachi_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_fibonachi_values.get(key), "-1.0")

		elif self.main_screen_coding_mode_button_data == "numerical":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_numerical_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_numerical_values.get(key), "-1.0")

	def __convert__(self):
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_pdf_to_word_converter_button.grid(column=1, row=0)
		self.main_screen_word_to_pdf_converter_button.grid(column=2, row=0)
		self.main_screen_pdf_to_txt_converter_button.grid(column=3, row=0)
		self.main_screen_txt_to_pdf_converter_button.grid(column=4, row=0)
		self.main_screen_word_to_txt_converter_button.grid(column=5, row=0)
		self.main_screen_txt_to_word_converter_button.grid(column=6, row=0)

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

	def __pdf_to_docx__(self):
		self.file = filedialog.askopenfilename(title="convert pdf file", filetypes=[("Pdf file (*.pdf)", "*.pdf")], defaultextension=[("Pdf file (*.pdf)", "*.pdf")])
		self.converted_file = Document(self.file)
		self.converted_file.save(self.file + ".docx")
		
	def __docx_to_pdf__(self):
		self.file = filedialog.askopenfilename(title="convert docx file", filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")])
		self.converted_file = Document(self.file)
		self.converted_file.save(self.file + ".pdf")

	def __pdf_to_txt__(self):
		self.file = filedialog.askopenfilename(title="convert pdf file", filetypes=[("Pdf file (*.pdf)", "*.pdf")], defaultextension=[("Pdf file (*.pdf)", "*.pdf")])
		self.converted_file = Document(self.file)
		self.converted_file.save(self.file + ".txt")
		
	def __txt_to_pdf__(self):
		self.file = filedialog.askopenfilename(title="convert txt file", filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")])
		self.converted_file = Document(self.file)
		self.converted_file.save(self.file + ".pdf")

	def __docx_to_txt__(self):
		self.file = filedialog.askopenfilename(title="convert docx file", filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")])
		self.converted_file = Document(self.file)
		self.converted_file.save(self.file + ".txt")
		
	def __txt_to_docx__(self):
		self.file = filedialog.askopenfilename(title="convert txt file", filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")])
		self.converted_file = Document(self.file)
		self.converted_file.save(self.file + ".docx")
		
	def __settings__(self):
		self.unbind('<Control_L>' + '<Tab>' + '<t>')

		self.main_screen_taskbar_button.grid(column=0, row=0)
		self.main_screen_settings_text.place(x=3, y=21)
		self.main_screen_settings_language_text.place(x=3, y=102)
		self.main_screen_settings_language_option.place(x=18, y=142)
		self.main_screen_settings_autosave_text.place(x=3, y=182)
		self.main_screen_settings_autosave_switch.place(x=18, y=222)
		self.main_screen_settings_theme_mode_text.place(x=3, y=262)
		self.main_screen_settings_theme_mode_option.place(x=18, y=302)

		self.main_screen_frame.place_forget()
		self.main_screen_word_counter.place_forget()

		self.main_screen_taskbar_button.configure(text="↵", command=self.__quit_operation__)

		self.main_screen_taskbar.grid_forget()

	def __language_settings__(self, pickle):
		self.main_screen_settings_language_option_variable = self.main_screen_settings_language_option.get()
		if self.main_screen_settings_language_option_variable == "Српски":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)
				
			showwarning(title="Пажња", message="Рестартуј програм")

		elif self.main_screen_settings_language_option_variable == "English":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)

			showwarning(title="Warning", message="Restart programm")

		elif self.main_screen_settings_language_option_variable == "Русский":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)

			showwarning(title="Внимание", message="Перезагрузите программу")

	def __autosave_settings__(self):
		self.main_screen_settings_autosave_switch_value = self.main_screen_settings_autosave_value.get()
		if self.main_screen_settings_autosave_switch_value == "on":
			with open("my_diary_autosave_settings.pickle", "wb") as self.autosave_data:
				dump(self.main_screen_settings_autosave_switch_value, self.autosave_data)

				if self.language_data == "Српски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")
		
		elif self.main_screen_settings_autosave_switch_value == "off":
			with open("my_diary_autosave_settings.pickle", "wb") as self.autosave_data:
				dump(self.main_screen_settings_autosave_switch_value, self.autosave_data)

				if self.language_data == "Српски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

	def __text_autosave__(self, event):
		self.main_screen_frame_textbox_text_data = self.main_screen_frame_textbox.get("1.0", END)

		with open('my_diary_saved_text.pickle', 'wb') as self.text_data:
			dump(self.main_screen_frame_textbox_text_data, self.text_data)

	def __theme_settings__(self, pickle):
		self.main_screen_settings_theme_mode_option_data = self.main_screen_settings_theme_mode_option.get()

		if self.main_screen_settings_theme_mode_option_data == "System":
			with open("my_diary_theme_settings.pickle", "wb") as self.theme_data:
				dump(self.main_screen_settings_theme_mode_option_data, self.theme_data)

				if self.language_data == "Српски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

		elif self.main_screen_settings_theme_mode_option_data == "Dark":
			with open("my_diary_theme_settings.pickle", "wb") as self.theme_data:
				dump(self.main_screen_settings_theme_mode_option_data, self.theme_data)

				if self.language_data == "Српски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

		elif self.main_screen_settings_theme_mode_option_data == "Light":
			with open("my_diary_theme_settings.pickle", "wb") as self.theme_data:
				dump(self.main_screen_settings_theme_mode_option_data, self.theme_data)

				if self.language_data == "Српски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

	def __word_count__(self, event):
		self.main_screen_frame_textbox_data = self.main_screen_frame_textbox.get("0.0", END)
		self.main_screen_word_counter_data_variable.set(value=len(self.main_screen_frame_textbox_data) - 1)
		
	def __program_version__(self, event):
		if self.language_data == "Српски":
			showinfo(title="Version_info", message="версија: 3.5")

		elif self.language_data == "English":
			showinfo(title="Version_info", message="version: 3.5")

		elif self.language_data == "Русский":
			showinfo(title="Version_info", message="версия: 3.5")

	def __fullscreen__(self, event):
		self.main_screen_fullscreen_numbers += 1
		if self.main_screen_fullscreen_numbers % 2 == 0:
			self.attributes('-fullscreen', True)
			
			self.main_screen_frame.configure(height=840.55)
			self.main_screen_frame_textbox.configure(height=839)

			self.main_screen_coding_frame.configure(height=840.55)
			self.main_screen_coding_frame_textbox.configure(height=839)
		
		else:
			self.attributes('-fullscreen', False)

			self.main_screen_frame.configure(height=769)
			self.main_screen_frame_textbox.configure(height=767.5)

			self.main_screen_coding_frame.configure(height=769)
			self.main_screen_coding_frame_textbox.configure(height=767.5)

	def __open_terminal__(self, event):
		self.terminal_frame = Terminal(master=self)

	def __html_script__(self, event):
		self.main_screen_frame_textbox.insert("0.0", "<!DOCTYPE html>" + "\n" + "\n" + "<html lang='en'>" + "\n" + "<head>" + "\n" + "	<meta charset='utf-8' />" + "\n" + "	<title></title>" + "\n" + "</head>" + "\n" + "<body>" + "\n" + "\n" + "</body>" + "\n" + "</html>")

	def __right_click_menu__(self, event):
		try:
			self.main_screen_right_click_menu.tk_popup(event.x_root, event.y_root)

		finally:
			self.main_screen_right_click_menu.grab_release()

	def __copy__(self):
		self.main_screen_frame_textbox_text_data = self.main_screen_frame_textbox.selection_get()
		return self.main_screen_frame_textbox_text_data

	def __paste__(self):
		self.main_screen_frame_textbox.insert(self.main_screen_frame_textbox.index("insert"), self.main_screen_frame_textbox_text_data)

	def __cut__(self):
		self.main_screen_frame_textbox.delete(self.main_screen_frame_textbox.index("sel.first"), self.main_screen_frame_textbox.index("sel.last"))	  																																			   

	def __exit__(self):
		if self.language_data == "Српски":
			self.main_screen_exit = askyesno(title="излаз", message="желите да изађете?")
			if self.main_screen_exit: sys.exit()

		elif self.language_data == "English":
			self.main_screen_exit = askyesno(title="exit", message="would you like to exit?")
			if self.main_screen_exit: sys.exit()

		elif self.language_data == "Русский":
			self.main_screen_exit = askyesno(title="выход", message="желайте выйти?")
			if self.main_screen_exit: sys.exit()

class Program_taskbar(CTkFrame):
	def __init__(self, master=any, width=300, height=791, border_width=1, corner_radius=5, color="green") -> CTkFrame:
		CTkFrame.__init__(self, master, width, height, corner_radius, border_width, color)

class Terminal(CTkToplevel):
	def __init__(self, master, *args, **kwargs) -> CTkToplevel:
		CTkToplevel.__init__(self, master, *args, **kwargs)

		self.geometry("374x375")
		self.attributes("-toolwindow", True)
		self.resizable(False, False)
		self.title("My Diary command prompt")

		self.terminal_frame = CTkFrame(master=self, height=300, width=300, border_width=1, border_color="white")
		self.terminal_frame.place(x=0, y=0)

		self.terminal_textbox = CTkTextbox(master=self.terminal_frame, height=264, width=296, fg_color="black", text_color="green")
		self.terminal_textbox.place(x=1.5, y=1)

		self.terminal_textbox.insert("0.0", ">>>")
		self.terminal_textbox.configure(state="disabled")

		self.terminal_entry = CTkEntry(master=self.terminal_frame, height=30, width=296, fg_color="black", text_color="green")
		self.terminal_entry.place(x=1.5, y=267)

		self.terminal_entry.bind('<Return>', self.__terminal_action__)

	def __terminal_action__(self, event):
		self.terminal_entry_data = self.terminal_entry.get()

		self.terminal_textbox.configure(state="normal")
		if self.terminal_entry_data == "clear text":
			program.__clear_text__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open taskbar":
			program.__open_taskbar__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "close taskbar":
			program.__close_taskbar__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "quit operation":
			program.__quit_operation__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)
		
		elif self.terminal_entry_data == "save text as txt":
			program.__save_text_as_txt__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as docx":
			program.__save_text_as_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as python":
			program.__save_text_as_py__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as java":
			program.__save_text_as_java__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as kotlin":
			program.__save_text_as_kotlin__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as html":
			program.__save_text_as_html__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open txt file":
			program.__open_file_txt__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open docx file":
			program.__open_file_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open python script":
			program.__open_file_py__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open html script":
			program.__open_file_html__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert pdf to docx":
			program.__pdf_to_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert docx to pdf":
			program.__docx_to_pdf__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "programm version":
			program.__programm_version__(event)

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "html script":
			program.__html_script__(event)

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "help":
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, "commands:" + "\n" + "1) save text as ..." + "\n" + "2) open ... file / script" + "\n" + "3) clear text" + "\n" + "4) convert ... to ..." + "\n" + "5) exit program / terminal" + "\n" + "6) open/close taskbar" + "\n" + "7) quit operation" + "\n" + "8) html script" + "\n" + "9) programm version" + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "exit terminal":
			self.destroy()

		elif self.terminal_entry_data == "exit program":
			program.__exit__()
		
		else:
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + " - command doesn't exist" + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)
			
if __name__ == "__main__":
	program = Program()
	program.mainloop()