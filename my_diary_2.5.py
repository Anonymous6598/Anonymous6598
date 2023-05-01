from customtkinter import * 
from tkinter import * 
from pickle import *
from tkinter.messagebox import *
from aspose.words import *
from tkinter import filedialog
from docx.shared import RGBColor, Pt
import os, sys, docx

class Programm(CTk):
	def __init__(self, *args, **kwargs) -> CTk:
		CTk.__init__(self, *args, **kwargs)

		set_default_color_theme("green")
		set_widget_scaling(1.251)
		deactivate_automatic_dpi_awareness()

		self.title("My diary")
		self.geometry("1920x1080")
		self.iconbitmap("my diary icon.ico")

		self.bind('<F1>', self.__programm_version__)  
		self.bind('<Control_L>' + '<f>', self.__fullscreen__) 
		self.bind('<Escape>' + '<e>', self.__exit__)
		self.bind('<Alt_L>' + '<t>', self.__open_terminal__)

		self.main_screen_fullscreen_numbers = 1

		with open("my_diary_settings.pickle", "rb") as self.data:
			self.new_data = load(self.data)

		with open('my_diary_saved_text.pickle', "rb") as self.text_data:
			self.new_text_data = load(self.text_data)

		with open("my_diary_autosave_settings.pickle", "rb") as self.autosave_data:
			self.new_autosaved_data = load(self.autosave_data)

		with open("my_diary_theme_settings.pickle", "rb") as self.theme_data:
			self.new_theme_data = load(self.theme_data)
			
		self.main_screen_undo_button = CTkButton(master=self, text="⟲", height=20, width=10, corner_radius=0, fg_color="green", command=lambda: self.main_screen_frame_textbox.edit_undo())
		self.main_screen_undo_button.grid(column=0, row=0)

		self.main_screen_redo_button = CTkButton(master=self, text="⟳", height=20, width=10, corner_radius=0, fg_color="green", command=lambda: self.main_screen_frame_textbox.edit_redo())
		self.main_screen_redo_button.grid(column=1, row=0)

		self.main_screen_save_button = CTkButton(master=self, text="сачувај текст", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save__)
		self.main_screen_save_button.grid(column=2, row=0)

		self.main_screen_save_as_txt_button = CTkButton(master=self, text="сачувај као txt", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_txt__)

		self.main_screen_save_as_docx_button = CTkButton(master=self, text="сачувај као docx", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_docx__)

		self.main_screen_save_as_py_button = CTkButton(master=self, text="сачувај као python", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_py__)

		self.main_screen_save_as_java_button = CTkButton(master=self, text="сачувај као java", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_java__)

		self.main_screen_save_as_kotlin_button = CTkButton(master=self, text="сачувај као kotlin", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_kotlin__)

		self.main_screen_save_as_html_button = CTkButton(master=self, text="сачувај као html", height=20, width=25, corner_radius=0, fg_color="green", command=self.__save_text_as_html__)

		self.main_screen_save_button_clicks = 1

		self.main_screen_clear_button = CTkButton(master=self, text="обриши текст", height=20, width=25, corner_radius=0, fg_color="green", command=lambda: self.main_screen_frame_textbox.delete("1.0", END))
		self.main_screen_clear_button.grid(column=3, row=0)

		self.main_screen_edit_text_button = CTkButton(master=self, text="уреди текст", height=20, width=25, corner_radius=0, fg_color="green", command=self.__edit_text_)
		self.main_screen_edit_text_button.grid(column=4, row=0)

		self.main_screen_edit_font_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["System", "Terminal", "Ubuntu", "Script", "Roman", "Modern", "MS Serif"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", dropdown_hover_color="white", text_color="white", dropdown_text_color="white", command=self.__edit_text_font_and_size__)

		self.main_screen_edit_size_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["1", "2", "4", "5", "6", "8", "11", "12", "13", "14", "16", "18", "20", "22", "24", "26", "28", "30", "32", "34", "36", "38", "40", "42", "44", "46", "48", "50", "60", "70", "80", "90", "100"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", dropdown_hover_color="white", text_color="white", dropdown_text_color="white", command=self.__edit_text_font_and_size__)

		self.main_screen_edit_color_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["black", "white", "red", "green", "blue"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", dropdown_hover_color="white", text_color="white", dropdown_text_color="white", command=self.__edit_text_color__)

		self.main_screen_edit_clicks = 1

		self.main_screen_open_file_button = CTkButton(master=self, text="отвори фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file__)
		self.main_screen_open_file_button.grid(column=5, row=0)
	
		self.main_screen_open_file_txt_button = CTkButton(master=self, text="отвори txt фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_txt__)
	
		self.main_screen_open_file_docx_button = CTkButton(master=self, text="отвори docx фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_docx__)
	
		self.main_screen_open_file_py_button = CTkButton(master=self, text="отвори python фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_py__)

		self.main_screen_open_file_html_button = CTkButton(master=self, text="отвори html фајл", height=20, width=25, corner_radius=0, fg_color="green", command=self.__open_file_html__)

		self.main_screen_open_file_clicks = 1

		self.main_screen_coding_button = CTkButton(master=self, text="кодирање", height=20, width=25, corner_radius=0, fg_color="green", command=self.__coding__)
		self.main_screen_coding_button.grid(column=6, row=0)

		self.main_screen_coding_binar_values = {"a": "0", "b": "1", "c": "10", "d": "11", "e": "100", "f": "101", "g": "110", "h": "111", "i": "1000", "j": "1001", "k": "1010", "l": "1011", "m": "1100", "n": "1101", "o": "1110", "p": "1111", "q": "10000", "r": "10001", "s": "10010", "t": "10011", "u": "10100", "v": "10101", "w": "10111", "x": "11000", "y": "11001", "z": "11011", " ": " ", ".": "."}

		self.main_screen_coding_ceaser_values = {"a": "d", "b": "e", "c": "f", "d": "g", "e": "h", "f": "i", "g": "j", "h": "k", "i": "l", "j": "m", "k": "n", "l": "o", "m": "p", "n": "q", "o": "r", "p": "s", "q": "t", "r": "u", "s": "v", "t": "w", "u": "x", "v": "y", "w": "z", "x": "a", "y": "b", "z": "c", " ": " ", ".": "."}
		
		self.main_screen_coding_fibonachi_values = {"a": "0", "b": "1", "c": "2", "d": "3", "e": "5", "f": "8", "g": "13", "h": "21", "i": "34", "j": "55", "k": "89", "l": "144", "m": "233", "n": "377", "o": "610", "p": "987", "q": "1597", "r": "2584", "s": "4181", "t": "6765", "u": "10946", "v": "17711", "w": "28657", "x": "46368", "y": "75025", "z": "121393", " ": " ", ".": "."}

		self.main_screen_coding_numerical_values = {"a": "1", "b": "2", "c" : "3", "d": "4", "e": "5", "f": "6", "g": "7", "h": "8", "i": "9", "j": "10", "k": "11", "l": "12", "m": "13", "q": "14", "r": "15", "s": "16", "t": "17", "u": "18", "v": "19", "w": "20", "x": "21", "y": "22", "z": "23", " ": " ", ".": "."}

		self.main_screen_coding_mode_button = CTkComboBox(master=self, height=20, width=75, corner_radius=0, values=["binar", "ceaser", "fibonachi", "numerical"], fg_color="green", border_color="green", button_color="green", button_hover_color="green", dropdown_fg_color="green", text_color="white", dropdown_hover_color="white", dropdown_text_color="white")

		self.main_screen_coding_function_button = CTkButton(master=self, text="кодирај", height=20, width=25, corner_radius=0, fg_color="green", command=self.__message_encoding__)

		self.main_screen_coding_clicks = 1

		self.main_screen_pdf_converter_button = CTkButton(master=self, text="Pdf претварач", height=20, width=25, corner_radius=0, fg_color="green", command=self.__pdf_covector__)
		self.main_screen_pdf_converter_button.grid(column=7, row=0)

		self.main_screen_pdf_to_word_converter_button = CTkButton(master=self, text="из Pdf у docx", height=20, width=25, corner_radius=0, fg_color="green", command=self.__pdf_to_docx__)

		self.main_screen_word_to_pdf_converter_button = CTkButton(master=self, text="из docx у Pdf", height=20, width=25, corner_radius=0, fg_color="green", command=self.__docx_to_pdf__)

		self.main_screen_pdf_clicks = 1

		self.main_screen_settings_button = CTkButton(master=self, text="подешавања", height=20, width=25, corner_radius=0, fg_color="green", command=self.__settings__)
		self.main_screen_settings_button.grid(column=8, row=0)

		self.main_screen_settings_text = CTkLabel(master=self, text="подешавања", font=("Roboto Bold", 72))

		self.main_screen_settings_language_text = CTkLabel(master=self, text="језици", font=("Roboto Bold", 36))

		self.main_screen_settings_language_option = CTkSegmentedButton(master=self, values=["Србски", "English", "Русский", "Deutch", "Español"], selected_color="green", command=self.__language_settings__)

		self.language_data = str(self.new_data)

		self.main_screen_settings_language_option.set(self.language_data)
		
		if self.language_data == "Србски":
			self.main_screen_save_button.configure(text="сачувај текст")
			self.main_screen_save_as_txt_button.configure(text="сачувај као txt")
			self.main_screen_save_as_docx_button.configure(text="сачувај као docx")
			self.main_screen_save_as_py_button.configure(text="сачувај као python")
			self.main_screen_save_as_java_button.configure(text="сачувај као java")
			self.main_screen_save_as_kotlin_button.configure(text="сачувај као kotlin")
			self.main_screen_save_as_html_button.configure(text="сачувај као html")

			self.main_screen_clear_button.configure(text="обриши текст")
			self.main_screen_edit_text_button.configure(text="уреди текст")
			
			self.main_screen_open_file_button.configure(text="отвори фајл")
			self.main_screen_open_file_txt_button.configure(text="отвори txt фајл")
			self.main_screen_open_file_docx_button.configure(text="отвори docx фајл")
			self.main_screen_open_file_py_button.configure(text="отвори python фајл")
			self.main_screen_open_file_html_button.configure(text="отвори html фајл")

			self.main_screen_coding_button.configure(text="кодирање")
			self.main_screen_coding_function_button.configure(text="кодирај")

			self.main_screen_pdf_converter_button.configure(text="Pdf претварач")
			self.main_screen_pdf_to_word_converter_button.configure(text="из Pdf у docx")
			self.main_screen_word_to_pdf_converter_button.configure(text="из docx у Pdf")

			self.main_screen_settings_button.configure(text="подешавања")
			self.main_screen_settings_text.configure(text="подешавања")
			self.main_screen_settings_language_text.configure(text="језици")

		elif self.language_data == "English":
			self.main_screen_save_button.configure(text="save text")
			self.main_screen_save_as_txt_button.configure(text="save as txt")
			self.main_screen_save_as_docx_button.configure(text="save as docx")
			self.main_screen_save_as_py_button.configure(text="save as python")
			self.main_screen_save_as_java_button.configure(text="save as java")
			self.main_screen_save_as_kotlin_button.configure(text="save as kotlin")
			self.main_screen_save_as_html_button.configure(text="save as html")

			self.main_screen_clear_button.configure(text="clear text")
			self.main_screen_edit_text_button.configure(text="edit text")
			
			self.main_screen_open_file_button.configure(text="open file")
			self.main_screen_open_file_txt_button.configure(text="open txt file")
			self.main_screen_open_file_docx_button.configure(text="open docx file")
			self.main_screen_open_file_py_button.configure(text="open python file")
			self.main_screen_open_file_html_button.configure(text="open html file")

			self.main_screen_coding_button.configure(text="coding")
			self.main_screen_coding_function_button.configure(text="code")

			self.main_screen_pdf_converter_button.configure(text="Pdf converterer")
			self.main_screen_pdf_to_word_converter_button.configure(text="from Pdf to docx")
			self.main_screen_word_to_pdf_converter_button.configure(text="from docx to Pdf")

			self.main_screen_settings_button.configure(text="settings")
			self.main_screen_settings_text.configure(text="settings")
			self.main_screen_settings_language_text.configure(text="languages")

		elif self.language_data == "Русский":
			self.main_screen_save_button.configure(text="сохранить текст")
			self.main_screen_save_as_txt_button.configure(text="сохранить как текст")
			self.main_screen_save_as_docx_button.configure(text="сохранить как ворд")
			self.main_screen_save_as_py_button.configure(text="сохранить как python")
			self.main_screen_save_as_java_button.configure(text="сохранить как java")
			self.main_screen_save_as_kotlin_button.configure(text="сохранить как kotlin")
			self.main_screen_save_as_html_button.configure(text="сохранить как html")

			self.main_screen_clear_button.configure(text="удалить текст")
			self.main_screen_edit_text_button.configure(text="редактировать текст")
			
			self.main_screen_open_file_button.configure(text="открыть файл")
			self.main_screen_open_file_txt_button.configure(text="открыть текстовый файл")
			self.main_screen_open_file_docx_button.configure(text="открыть ворд файл")
			self.main_screen_open_file_py_button.configure(text="открыть python файл")
			self.main_screen_open_file_html_button.configure(text="открыть html файл")
			
			self.main_screen_coding_button.configure(text="кодировка")
			self.main_screen_coding_function_button.configure(text="кодировать")

			self.main_screen_pdf_converter_button.configure(text="конвертер Пдф")
			self.main_screen_pdf_to_word_converter_button.configure(text="из Пдф в ворд")
			self.main_screen_word_to_pdf_converter_button.configure(text="из ворд в Пдф")

			self.main_screen_settings_button.configure(text="настройки")
			self.main_screen_settings_text.configure(text="настройки")
			self.main_screen_settings_language_text.configure(text="языки")

		elif self.language_data == "Deutch":
			self.main_screen_save_button.configure(text="Texte speichern")
			self.main_screen_save_as_txt_button.configure(text="als txt speichern")
			self.main_screen_save_as_docx_button.configure(text="als doсx speichern")
			self.main_screen_save_as_py_button.configure(text="als Python speichern")
			self.main_screen_save_as_java_button.configure(text="als Java speichern")
			self.main_screen_save_as_kotlin_button.configure(text="als Kotlin speichern")
			self.main_screen_save_as_html_button.configure(text="als Html speichern")

			self.main_screen_clear_button.configure(text="Texte löschen")
			self.main_screen_edit_text_button.configure(text="Text bearbeiten")
			
			self.main_screen_open_file_button.configure(text="Datei öffnen")
			self.main_screen_open_file_txt_button.configure(text="Textdatei öffnen")
			self.main_screen_open_file_docx_button.configure(text="Word-Datei öffnen")
			self.main_screen_open_file_py_button.configure(text="Python-Datei öffnen")
			self.main_screen_open_file_html_button.configure(text="Html-Datei öffnen")
			
			self.main_screen_coding_button.configure(text="Codierung")
			self.main_screen_coding_function_button.configure(text="kodieren")

			self.main_screen_pdf_converter_button.configure(text="Pdf konverter")
			self.main_screen_pdf_to_word_converter_button.configure(text="von Pdf zu docx")
			self.main_screen_word_to_pdf_converter_button.configure(text="von docx zu pdf")

			self.main_screen_settings_button.configure(text="Einstellungen")
			self.main_screen_settings_text.configure(text="Einstellungen")
			self.main_screen_settings_language_text.configure(text="Sprachen")

		elif self.language_data == "Español":
			self.main_screen_save_button.configure(text="guardar texto")
			self.main_screen_save_as_txt_button.configure(text="guardar como txt")
			self.main_screen_save_as_docx_button.configure(text="guardar como docx")
			self.main_screen_save_as_py_button.configure(text="guardar como python")
			self.main_screen_save_as_java_button.configure(text="guardar como java")
			self.main_screen_save_as_kotlin_button.configure(text="guardar como kotlin")
			self.main_screen_save_as_html_button.configure(text="guardar como kotlin")

			self.main_screen_clear_button.configure(text="borrar texto")
			self.main_screen_edit_text_button.configure(text="editar texto")
			
			self.main_screen_open_file_button.configure(text="abrir documento")
			self.main_screen_open_file_txt_button.configure(text="abrir archivo de texto")
			self.main_screen_open_file_docx_button.configure(text="abrir archivo de palabra")
			self.main_screen_open_file_py_button.configure(text="abrir archivo python")
			self.main_screen_open_file_html_button.configure(text="abrir archivo html")
			
			self.main_screen_coding_button.configure(text="codificacióng")
			self.main_screen_coding_function_button.configure(text="codificar")

			self.main_screen_pdf_converter_button.configure(text="Pdf konverter")
			self.main_screen_pdf_to_word_converter_button.configure(text="de Pdf a docx")
			self.main_screen_word_to_pdf_converter_button.configure(text="de docx a Pdf")

			self.main_screen_settings_button.configure(text="ajustes")
			self.main_screen_settings_text.configure(text="ajustes")
			self.main_screen_settings_language_text.configure(text="lenguas")

		self.main_screen_settings_autosave_text = CTkLabel(master=self, text="autosave", font=("Roboto Bold", 36))

		self.main_screen_settings_autosave_value = StringVar(value=self.new_autosaved_data)

		self.main_screen_settings_autosave_switch_variable = str(self.new_autosaved_data)

		self.main_screen_settings_autosave_switch = CTkSwitch(master=self, text="autosave", variable=self.main_screen_settings_autosave_value,onvalue="on", offvalue="off", progress_color="green", command=self.__autosave_settings__)
			
		self.main_screen_settings_theme_mode_text = CTkLabel(master=self, text="тема", font=("Roboto Bold", 36))

		self.main_screen_settings_theme_mode_option = CTkSegmentedButton(master=self, values=["System", "Dark", "Light"], selected_color="green", command=self.__theme_settings__)

		self.theme = str(self.new_theme_data)

		self.main_screen_settings_theme_mode_option.set(self.theme)

		if self.theme == "System":
			set_appearance_mode("system")

		elif self.theme == "Dark":
			set_appearance_mode("dark")

		elif self.theme == "Light":
			set_appearance_mode("light")

		self.main_screen_settings_button_clicks = 1

		self.main_screen_frame = CTkFrame(master=self, height=769, width=1535, border_width=1, border_color=("black", "white"), corner_radius=0)
		self.main_screen_frame.place(x=0, y=22.35)

		self.main_screen_frame_textbox = CTkTextbox(master=self.main_screen_frame, height=767.5, width=1533, corner_radius=0, undo=True)
		self.main_screen_frame_textbox.place(x=1, y=1)

		self.main_screen_frame_textbox.bind('<KeyRelease>', self.__word_count__)
		self.main_screen_frame_textbox.bind('<F2>', self.__html_script__)

		self.autosaved_text = str(self.new_text_data)

		self.main_screen_settings_autosave_switch_value = self.main_screen_settings_autosave_value.get()

		if self.main_screen_settings_autosave_switch_variable == "on":
			self.main_screen_frame_textbox.insert("1.0", self.autosaved_text)

			self.bind('<Return>', self.__text_autosave__)
			self.bind('<Control_L>' + '<s>', self.__text_autosave__)

		elif self.main_screen_settings_autosave_switch_variable == "off":
			self.unbind('<Return>')
			self.unbind('<Control_L>' + '<s>')

		self.main_screen_word_counter_variable = 0

		self.main_screen_word_counter_data_variable = IntVar(value=self.main_screen_word_counter_variable)

		self.main_screen_word_counter = CTkLabel(master=self, width=25, height=20, corner_radius=0, textvariable=self.main_screen_word_counter_data_variable)
		self.main_screen_word_counter.place(x=1500, y=0)
				  
		self.main_screen_coding_frame = CTkFrame(master=self, height=769, width=807, border_width=1, border_color=("black", "white"), corner_radius=0)

		self.main_screen_coding_frame_textbox = CTkTextbox(master=self.main_screen_coding_frame, height=767.5, width=804, corner_radius=0)

	def __clear_text__(self):
		self.main_screen_frame_textbox.delete("1.0", END)

	def __save__(self):
		self.main_screen_save_button_clicks += 1
		if self.main_screen_save_button_clicks % 2 == 0:
			self.main_screen_save_button.grid(column=0, row=0)
			self.main_screen_save_as_txt_button.grid(column=1, row=0)
			self.main_screen_save_as_docx_button.grid(column=2, row=0)
			self.main_screen_save_as_py_button.grid(column=3, row=0)
			self.main_screen_save_as_java_button.grid(column=4, row=0)
			self.main_screen_save_as_kotlin_button.grid(column=5, row=0)
			self.main_screen_save_as_html_button.grid(column=6, row=0)

			self.main_screen_undo_button.grid_forget()
			self.main_screen_redo_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_coding_button.grid_forget()
			self.main_screen_pdf_converter_button.grid_forget()
			self.main_screen_settings_button.grid_forget()
			self.main_screen_word_counter.place_forget()

		else:
			self.main_screen_undo_button.grid(column=0, row=0)
			self.main_screen_redo_button.grid(column=1, row=0)
			self.main_screen_save_button.grid(column=2, row=0)
			self.main_screen_clear_button.grid(column=3, row=0)
			self.main_screen_edit_text_button.grid(column=4, row=0)
			self.main_screen_open_file_button.grid(column=5, row=0)
			self.main_screen_coding_button.grid(column=6, row=0)
			self.main_screen_pdf_converter_button.grid(column=7, row=0)
			self.main_screen_settings_button.grid(column=8, row=0)
			self.main_screen_word_counter.place(x=1500, y=0)

			self.main_screen_save_as_txt_button.grid_forget()
			self.main_screen_save_as_docx_button.grid_forget()
			self.main_screen_save_as_py_button.grid_forget()
			self.main_screen_save_as_java_button.grid_forget()
			self.main_screen_save_as_kotlin_button.grid_forget()
			self.main_screen_save_as_html_button.grid_forget()
			
	def __save_text_as_txt__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_docx__(self):
		try:
			self.file = docx.Document()
			self.file_data = self.main_screen_frame_textbox.get("1.0", END)
			self.file_run = self.file.add_paragraph().add_run(self.file_data)
			self.font = self.file_run.font
			self.font.name = str(self.main_screen_edit_font_button_data)
			self.font.size = Pt(int(self.main_screen_edit_size_button_data))
			if self.main_screen_edit_color_button_data == "black":
				self.font.color.rgb = RGBColor(0, 0, 0)

			elif self.main_screen_edit_color_button_data == "white":
				self.font.color.rgb = RGBColor(255, 255, 255)

			elif self.main_screen_edit_color_button_data == "red":
				self.font.color.rgb = RGBColor(250, 0, 0)

			elif self.main_screen_edit_color_button_data == "green":
				self.font.color.rgb = RGBColor(0, 255, 0)

			elif self.main_screen_edit_color_button_data == "blue":
				self.font.color.rgb = RGBColor(0, 0, 255)

			self.file.save(filedialog.asksaveasfilename(filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")]))
		
		except AttributeError:
			if self.language_data == "Србски":
				showerror(title="Грешка", message="Појавила се грешка. Нисам мого да сачувам фајл")

			elif self.language_data == "English":
				showerror(title="Error", message="An error has occured. Couldn't save the file")

			elif self.language_data == "Русский":
				showerror(title="Ошибка", message="Появилась ошибка. Не смог охранить файл")

			elif self.language_data == "Deutch":
				showwarning(title="Ahtung", message="Starten Sie das Programm neu")

			elif self.language_data == "Español":
				showwarning(title="Atención", message="Reiniciar el programa")

	def __save_text_as_py__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Python file (*.py)", "*.py")], defaultextension=[("Python file (*.py)", "*.py")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_java__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Java file (*.java)", "*.java")], defaultextension=[("Java file (*.java)", "*.java")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_kotlin__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("Kotlin file (*.kt)", "*.kt")], defaultextension=[("Kotlin file (*.kt)", "*.kt")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __save_text_as_html__(self):
		self.file = open(filedialog.asksaveasfilename(filetypes=[("html file (*.html)", "*.html")], defaultextension=[("Python file (*.html)", "*.html")]), "w")
		file_data = self.main_screen_frame_textbox.get("1.0", END)
		self.file.write(file_data)
		self.file.close()

	def __edit_text_(self):
		self.main_screen_edit_clicks += 1
		if self.main_screen_edit_clicks % 2 == 0:
			self.main_screen_edit_text_button.grid(column=0, row=0)
			self.main_screen_edit_font_button.grid(column=1, row=0)
			self.main_screen_edit_size_button.grid(column=2, row=0)
			self.main_screen_edit_color_button.grid(column=3, row=0)

			self.main_screen_undo_button.grid_forget()
			self.main_screen_redo_button.grid_forget()
			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_coding_button.grid_forget()
			self.main_screen_pdf_converter_button.grid_forget()
			self.main_screen_settings_button.grid_forget()
			self.main_screen_word_counter.place_forget()

		else:
			self.main_screen_undo_button.grid(column=0, row=0)
			self.main_screen_redo_button.grid(column=1, row=0)
			self.main_screen_save_button.grid(column=2, row=0)
			self.main_screen_clear_button.grid(column=3, row=0)
			self.main_screen_edit_text_button.grid(column=4, row=0)
			self.main_screen_open_file_button.grid(column=5, row=0)
			self.main_screen_coding_button.grid(column=6, row=0)
			self.main_screen_pdf_converter_button.grid(column=7, row=0)
			self.main_screen_settings_button.grid(column=8, row=0)
			self.main_screen_word_counter.place(x=1500, y=0)

			self.main_screen_edit_font_button.grid_forget()
			self.main_screen_edit_size_button.grid_forget()
			self.main_screen_edit_color_button.grid_forget()

	def __edit_text_font_and_size__(self, configure):
		self.main_screen_edit_font_button_data = self.main_screen_edit_font_button.get()
		self.main_screen_edit_size_button_data = self.main_screen_edit_size_button.get()
		self.main_screen_frame_textbox.configure(font=(self.main_screen_edit_font_button_data, float(self.main_screen_edit_size_button_data)))

	def __edit_text_color__(self, configure):
		self.main_screen_edit_color_button_data = self.main_screen_edit_color_button.get()
		self.main_screen_frame_textbox.configure(text_color=self.main_screen_edit_color_button_data)

	def __open_file__(self):
		self.main_screen_open_file_clicks += 1
		if self.main_screen_open_file_clicks % 2 == 0:
			self.main_screen_open_file_button.grid(column=0, row=0)
			self.main_screen_open_file_txt_button.grid(column=1, row=0)
			self.main_screen_open_file_docx_button.grid(column=2, row=0)
			self.main_screen_open_file_py_button.grid(column=3, row=0)
			self.main_screen_open_file_html_button.grid(column=4, row=0)

			self.main_screen_undo_button.grid_forget()
			self.main_screen_redo_button.grid_forget()
			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_coding_button.grid_forget()
			self.main_screen_pdf_converter_button.grid_forget()
			self.main_screen_settings_button.grid_forget()
			self.main_screen_word_counter.place_forget()

		else:
			self.main_screen_undo_button.grid(column=0, row=0)
			self.main_screen_redo_button.grid(column=1, row=0)
			self.main_screen_save_button.grid(column=2, row=0)
			self.main_screen_clear_button.grid(column=3, row=0)
			self.main_screen_edit_text_button.grid(column=4, row=0)
			self.main_screen_open_file_button.grid(column=5, row=0)
			self.main_screen_coding_button.grid(column=6, row=0)
			self.main_screen_pdf_converter_button.grid(column=7, row=0)
			self.main_screen_settings_button.grid(column=8, row=0)
			self.main_screen_word_counter.place(x=1500, y=0)

			self.main_screen_open_file_txt_button.grid_forget()
			self.main_screen_open_file_docx_button.grid_forget()
			self.main_screen_open_file_py_button.grid_forget()
			self.main_screen_open_file_html_button.grid_forget()

	def __open_file_txt__(self):
		self.openned_file = open(filedialog.askopenfilename(title="open txt file", filetypes=[("Text file (*.txt)", "*.txt")], defaultextension=[("Text file (*.txt)", "*.txt")]), "r+")
		self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

	def __open_file_docx__(self):
		self.openned_file_name = filedialog.askopenfilename(title="open docx file", filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")])
		self.openned_file = docx.Document(self.openned_file_name)
		self.openned_file_data = []
		for self.paragraphs in self.openned_file.paragraphs:
			self.openned_file_data.append(self.paragraphs.text)

		self.main_screen_frame_textbox.insert("1.0", "\n".join(self.openned_file_data))

	def __open_file_py__(self):
		self.openned_file = open(filedialog.askopenfilename(title="open python file", filetypes=[("Python file (*.py)", "*.py")], defaultextension=[("Python file (*.py)", "*.py")]), "r+", encoding="UTF-8")
		self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

	def __open_file_html__(self):
		self.openned_file = open(filedialog.askopenfilename(title="open html file", filetypes=[("html file (*.html)", "*.html")], defaultextension=[("html file (*.html)", "*.html")]), "r+", encoding="UTF-8")
		self.main_screen_frame_textbox.insert("1.0", self.openned_file.read())

	def __coding__(self):
		self.main_screen_coding_clicks += 1
		if self.main_screen_coding_clicks % 2 == 0:
			self.main_screen_frame.configure(width=715.5)
			self.main_screen_frame_textbox.configure(width=712.75)

			self.main_screen_coding_button.grid(column=0, row=0)
			self.main_screen_coding_frame.place(x=729, y=22)
			self.main_screen_coding_frame_textbox.place(x=1, y=1)
			self.main_screen_coding_mode_button.place(x=729, y=0)
			self.main_screen_coding_function_button.place(x=803.5, y=0)

			self.main_screen_undo_button.grid_forget()
			self.main_screen_redo_button.grid_forget()
			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_pdf_converter_button.grid_forget()
			self.main_screen_settings_button.grid_forget()
			self.main_screen_word_counter.place_forget()

		else:
			self.main_screen_frame.configure(width=1535)
			self.main_screen_frame_textbox.configure(width=1530)

			self.main_screen_undo_button.grid(column=0, row=0)
			self.main_screen_redo_button.grid(column=1, row=0)
			self.main_screen_save_button.grid(column=2, row=0)
			self.main_screen_clear_button.grid(column=3, row=0)
			self.main_screen_edit_text_button.grid(column=4, row=0)
			self.main_screen_open_file_button.grid(column=5, row=0)
			self.main_screen_coding_button.grid(column=6, row=0)
			self.main_screen_pdf_converter_button.grid(column=7, row=0)
			self.main_screen_settings_button.grid(column=8, row=0)
			self.main_screen_word_counter.place(x=1500, y=0)

			self.main_screen_coding_frame.place_forget()
			self.main_screen_coding_frame_textbox.place_forget()
			self.main_screen_coding_mode_button.place_forget()
			self.main_screen_coding_function_button.place_forget()

			self.main_screen_coding_frame_textbox.delete("1.0", END)

	def __message_encoding__(self):
		self.main_screen_coding_frame_textbox.delete("1.0", END)

		self.main_screen_coding_mode_button_data = self.main_screen_coding_mode_button.get()
		self.main_screen_frame_textbox_data = self.main_screen_frame_textbox.get("1.0", END)
		if self.main_screen_coding_mode_button_data == "binar":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_binar_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_binar_values.get(key), "-1.0")
		
		elif self.main_screen_coding_mode_button_data == "ceaser":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_ceaser_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_ceaser_values.get(key), "-1.0")

		elif self.main_screen_coding_mode_button_data == "fibonachi":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_fibonachi_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_fibonachi_values.get(key), "-1.0")

		elif self.main_screen_coding_mode_button_data == "numerical":
			for symbols in self.main_screen_frame_textbox_data:
				for key in self.main_screen_coding_numerical_values:
					if key == symbols:
						self.main_screen_coding_frame_textbox.insert(END, self.main_screen_coding_numerical_values.get(key), "-1.0")

	def __pdf_covector__(self):
		self.main_screen_pdf_clicks += 1
		if self.main_screen_pdf_clicks % 2 == 0:
			self.main_screen_pdf_converter_button.grid(column=0, row=0)
			self.main_screen_pdf_to_word_converter_button.grid(column=1, row=0)
			self.main_screen_word_to_pdf_converter_button.grid(column=2, row=0)

			self.main_screen_undo_button.grid_forget()
			self.main_screen_redo_button.grid_forget()
			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_coding_button.grid_forget()
			self.main_screen_settings_button.grid_forget()
			self.main_screen_word_counter.place_forget()

		else:
			self.main_screen_undo_button.grid(column=0, row=0)
			self.main_screen_redo_button.grid(column=1, row=0)
			self.main_screen_save_button.grid(column=2, row=0)
			self.main_screen_clear_button.grid(column=3, row=0)
			self.main_screen_edit_text_button.grid(column=4, row=0)
			self.main_screen_open_file_button.grid(column=5, row=0)
			self.main_screen_coding_button.grid(column=6, row=0)
			self.main_screen_pdf_converter_button.grid(column=7, row=0)
			self.main_screen_settings_button.grid(column=8, row=0)
			self.main_screen_word_counter.place(x=1500, y=0)

			self.main_screen_pdf_to_word_converter_button.grid_forget()
			self.main_screen_word_to_pdf_converter_button.grid_forget()

	def __pdf_to_docx__(self):
		self.file = filedialog.askopenfilename(title="convert pdf file", filetypes=[("Pdf file (*.pdf)", "*.pdf")], defaultextension=[("Pdf file (*.pdf)", "*.pdf")])
		self.converted_file = Document(self.file)
		self.converted_file.save(self.file + ".docx")
		
	def __docx_to_pdf__(self):
		self.file = filedialog.askopenfilename(title="convert docx file", filetypes=[("Word file (*.docx)", "*.docx")], defaultextension=[("Word file (*.docx)", "*.docx")])
		self.converted_file = Document(self.file)
		self.converted_file.save(self.file + ".pdf")
		
	def __settings__(self):
		self.main_screen_settings_button_clicks += 1
		if self.main_screen_settings_button_clicks % 2 == 0:
			self.main_screen_settings_button.grid(column=0, row=0)
			self.main_screen_settings_text.place(x=3, y=21)
			self.main_screen_settings_language_text.place(x=3, y=102)
			self.main_screen_settings_language_option.place(x=18, y=142)
			self.main_screen_settings_autosave_text.place(x=3, y=182)
			self.main_screen_settings_autosave_switch.place(x=18, y=222)
			self.main_screen_settings_theme_mode_text.place(x=3, y=262)
			self.main_screen_settings_theme_mode_option.place(x=18, y=302)

			self.main_screen_undo_button.grid_forget()
			self.main_screen_redo_button.grid_forget()
			self.main_screen_save_button.grid_forget()
			self.main_screen_clear_button.grid_forget()
			self.main_screen_edit_text_button.grid_forget()
			self.main_screen_open_file_button.grid_forget()
			self.main_screen_coding_button.grid_forget()
			self.main_screen_pdf_converter_button.grid_forget()
			self.main_screen_frame.place_forget()
			self.main_screen_word_counter.place_forget()
	   
		else:
			self.main_screen_undo_button.grid(column=0, row=0)
			self.main_screen_redo_button.grid(column=1, row=0)
			self.main_screen_save_button.grid(column=2, row=0)
			self.main_screen_clear_button.grid(column=3, row=0)
			self.main_screen_edit_text_button.grid(column=4, row=0)
			self.main_screen_open_file_button.grid(column=5, row=0)
			self.main_screen_coding_button.grid(column=6, row=0)
			self.main_screen_pdf_converter_button.grid(column=7, row=0)
			self.main_screen_settings_button.grid(column=8, row=0)
			self.main_screen_frame.place(x=0, y=22.35)
			self.main_screen_word_counter.place(x=1500, y=0)

			self.main_screen_settings_text.place_forget()
			self.main_screen_settings_language_text.place_forget()
			self.main_screen_settings_language_option.place_forget()						          
			self.main_screen_settings_autosave_text.place_forget()
			self.main_screen_settings_autosave_switch.place_forget()
			self.main_screen_settings_theme_mode_text.place_forget()
			self.main_screen_settings_theme_mode_option.place_forget()

	def __language_settings__(self, pickle):
		self.main_screen_settings_language_option_variable = self.main_screen_settings_language_option.get()
		if self.main_screen_settings_language_option_variable == "Србски":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)
				
			showwarning(title="Пажња", message="Рестартуј програм")

		elif self.main_screen_settings_language_option_variable == "English":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)

			showwarning(title="Warning", message="Restart programm")

		elif self.main_screen_settings_language_option_variable == "Русский":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)

			showwarning(title="Внимание", message="Перезагрузите программу")

		elif self.main_screen_settings_language_option_variable == "Deutch":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)

			showwarning(title="Ahtung", message="Starten Sie das Programm neu")

		elif self.main_screen_settings_language_option_variable == "Español":
			with open("my_diary_settings.pickle", "wb") as self.data:
				dump(self.main_screen_settings_language_option_variable, self.data)

			showwarning(title="Atención", message="Reiniciar el programa")

	def __autosave_settings__(self):
		self.main_screen_settings_autosave_switch_value = self.main_screen_settings_autosave_value.get()
		if self.main_screen_settings_autosave_switch_value == "on":
			with open("my_diary_autosave_settings.pickle", "wb") as self.autosave_data:
				dump(self.main_screen_settings_autosave_switch_value, self.autosave_data)

				if self.language_data == "Србски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

				elif self.language_data == "Deutch":
					showwarning(title="Ahtung", message="Starten Sie das Programm neu")

				elif self.language_data == "Español":
					showwarning(title="Atención", message="Reiniciar el programa")
		
		elif self.main_screen_settings_autosave_switch_value == "off":
			with open("my_diary_autosave_settings.pickle", "wb") as self.autosave_data:
				dump(self.main_screen_settings_autosave_switch_value, self.autosave_data)

				if self.language_data == "Србски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

				elif self.language_data == "Deutch":
					showwarning(title="Ahtung", message="Starten Sie das Programm neu")

				elif self.language_data == "Español":
					showwarning(title="Atención", message="Reiniciar el programa")

	def __text_autosave__(self, event):
		self.main_screen_frame_textbox_text_data = self.main_screen_frame_textbox.get("1.0", END)

		with open('my_diary_saved_text.pickle', 'wb') as self.text_data:
			dump(self.main_screen_frame_textbox_text_data, self.text_data)

	def __theme_settings__(self, pickle):
		self.main_screen_settings_theme_mode_option_data = self.main_screen_settings_theme_mode_option.get()

		if self.main_screen_settings_theme_mode_option_data == "System":
			with open("my_diary_theme_settings.pickle", "wb") as self.theme_data:
				dump(self.main_screen_settings_theme_mode_option_data, self.theme_data)

				if self.language_data == "Србски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

				elif self.language_data == "Deutch":
					showwarning(title="Ahtung", message="Starten Sie das Programm neu")

				elif self.language_data == "Español":
					showwarning(title="Atención", message="Reiniciar el programa")

		if self.main_screen_settings_theme_mode_option_data == "Dark":
			with open("my_diary_theme_settings.pickle", "wb") as self.theme_data:
				dump(self.main_screen_settings_theme_mode_option_data, self.theme_data)

				if self.language_data == "Србски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

				elif self.language_data == "Deutch":
					showwarning(title="Ahtung", message="Starten Sie das Programm neu")

				elif self.language_data == "Español":
					showwarning(title="Atención", message="Reiniciar el programa")

		if self.main_screen_settings_theme_mode_option_data == "Light":
			with open("my_diary_theme_settings.pickle", "wb") as self.theme_data:
				dump(self.main_screen_settings_theme_mode_option_data, self.theme_data)

				if self.language_data == "Србски":
					showwarning(title="Пажња", message="Рестартуј програм")

				elif self.language_data == "English":
					showwarning(title="Warning", message="Restart programm")

				elif self.language_data == "Русский":
					showwarning(title="Внимание", message="Перезагрузите программу")

				elif self.language_data == "Deutch":
					showwarning(title="Ahtung", message="Starten Sie das Programm neu")

				elif self.language_data == "Español":
					showwarning(title="Atención", message="Reiniciar el programa")

	def __word_count__(self, event):
		self.main_screen_frame_textbox_data = self.main_screen_frame_textbox.get("0.0", END)
		self.main_screen_word_counter_data_variable.set(value=len(self.main_screen_frame_textbox_data) - 1)
		
	def __programm_version__(self, event):
		if self.language_data == "Србски":
			showinfo(title="Version_info", message="версија: 2.5")

		elif self.language_data == "English":
			showinfo(title="Version_info", message="version: 2.5")

		elif self.language_data == "Русский":
			showinfo(title="Version_info", message="версия: 2.5")

		elif self.language_data == "Deutch":
			showinfo(title="Version_info", message="Ausführung: 2.5")

		elif self.language_data == "Español":
			showinfo(title="Version_info", message="versión: 2.5")

	def __fullscreen__(self, event):
		self.main_screen_fullscreen_numbers += 1
		if self.main_screen_fullscreen_numbers % 2 == 0:
			self.attributes('-fullscreen', True)
			
			self.main_screen_frame.configure(height=840.55)
			self.main_screen_frame_textbox.configure(height=839)

			self.main_screen_coding_frame.configure(height=840.55)
			self.main_screen_coding_frame_textbox.configure(height=839)
		
		else:
			self.attributes('-fullscreen', False)

			self.main_screen_frame.configure(height=769)
			self.main_screen_frame_textbox.configure(height=767.5)

			self.main_screen_coding_frame.configure(height=769)
			self.main_screen_coding_frame_textbox.configure(height=767.5)

	def __open_terminal__(self, event):
		self.terminal_frame = Terminal(master=self)

	def __html_script__(self, event):
		self.main_screen_frame_textbox.insert("0.0", "<!DOCTYPE html>" + "\n" + "\n" + "<html lang='en'>" + "\n" + "<head>" + "\n" + "	<meta charset='utf-8' />" + "\n" + "	<title></title>" + "\n" + "</head>" + "\n" + "<body>" + "\n" + "\n" + "</body>" + "\n" + "</html>")

	def __exit__(self, event):
		if self.language_data == "Србски":
			self.main_screen_exit = askyesno(title="излаз", message="желите да изађете?")
			if self.main_screen_exit: sys.exit()

		elif self.language_data == "English":
			self.main_screen_exit = askyesno(title="exit", message="would you like to exit?")
			if self.main_screen_exit: sys.exit()

		elif self.language_data == "Русский":
			self.main_screen_exit = askyesno(title="выход", message="желайте выйти?")
			if self.main_screen_exit: sys.exit()

		elif self.language_data == "Deutch":
			self.main_screen_exit = askyesno(title="ausgehen", message="ausgehen wollen?")
			if self.main_screen_exit: sys.exit()

		elif self.language_data == "Español":
			self.main_screen_exit = askyesno(title="salir", message="quiero salir?")
			if self.main_screen_exit: sys.exit()

class Terminal(CTkToplevel):
	def __init__(self, master, *args, **kwargs) -> CTkToplevel:
		CTkToplevel.__init__(self, master, *args, **kwargs)

		self.geometry("374x375")
		self.attributes("-toolwindow", True)
		self.resizable(False, False)
		self.title("My Diary terminal shell")

		self.terminal_frame = CTkFrame(master=self, height=300, width=300, border_width=1, border_color="white")
		self.terminal_frame.place(x=0, y=0)

		self.terminal_textbox = CTkTextbox(master=self.terminal_frame, height=264, width=296, fg_color="black", text_color="green")
		self.terminal_textbox.place(x=1.5, y=1)

		self.terminal_textbox.insert("0.0", ">>>")
		self.terminal_textbox.configure(state="disabled")

		self.terminal_entry = CTkEntry(master=self.terminal_frame, height=30, width=296, fg_color="black", text_color="green")
		self.terminal_entry.place(x=1.5, y=267)

		self.terminal_entry.bind('<Return>', self.__terminal_action__)

	def __terminal_action__(self, event):
		self.terminal_entry_data = self.terminal_entry.get()

		self.terminal_textbox.configure(state="normal")
		if self.terminal_entry_data == "clear text":
			programm.__clear_text__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)
		
		elif self.terminal_entry_data == "save text as txt":
			programm.__save_text_as_txt__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as docx":
			programm.__save_text_as_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as python":
			programm.__save_text_as_py__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as java":
			programm.__save_text_as_java__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as kotlin":
			programm.__save_text_as_kotlin__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "save text as html":
			programm.__save_text_as_html__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open txt file":
			programm.__open_file_txt__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open docx file":
			programm.__open_file_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open python script":
			programm.__open_file_py__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "open html script":
			programm.__open_file_html__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert pdf to docx":
			programm.__pdf_to_docx__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "convert docx to pdf":
			programm.__docx_to_pdf__()

			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "help":
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, "commands:" + "\n" + "1) save text as ..." + "\n" + "2) open ... file / script" + "\n" + "3) clear text" + "\n" + "4) convert ... to ..." + "\n" + "5) exit programm / terminal" + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)

		elif self.terminal_entry_data == "exit terminal":
			self.destroy()

		elif self.terminal_entry_data == "exit programm":
			sys.exit()
		
		else:
			self.terminal_textbox.configure(state="normal")
			self.terminal_textbox.insert(END, self.terminal_entry_data + " command doesn't exist" + "\n", "-1.0")
			self.terminal_textbox.insert(END, ">>>", "-1.0")
			self.terminal_textbox.configure(state="disabled")
			self.terminal_entry.delete("-1", END)
			
if __name__ == "__main__":
	programm = Programm()
	programm.mainloop()